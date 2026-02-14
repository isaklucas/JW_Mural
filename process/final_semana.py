"""
Módulo para geração de documentos de Reunião de Final de Semana.
Extrai títulos da Sentinela do wol.jw.org, solicita dados do usuário e gera documento Word.
"""
import subprocess
import datetime
import json
import logging
import os
from docx import Document
import process.webscrapper as webscrapper
from database import db_ops
import tkinter as tk
import ttkbootstrap as ttk
from ttkbootstrap.constants import *

logger = logging.getLogger(__name__)


class FinalSemana:
    def gerar_final_semana(url, nome_arquivo, idioma, usar_selecao_automatica_leitor, usar_selecao_automatica_presidente=False, parent=None):
        """
        Orquestra o fluxo completo de geração de reunião final de semana (9 semanas, dois documentos).
        parent: janela pai para o modal (obrigatório quando chamado de thread - evita RuntimeError).
        """
        try:
            semanas = FinalSemana.buscar_titulos_meetings(url, 9)
            if not semanas:
                raise ValueError("Não foi possível extrair os títulos da Sentinela. Verifique a URL.")
            
            if usar_selecao_automatica_presidente:
                FinalSemana.selecionar_presidente_automaticamente(semanas)
            if usar_selecao_automatica_leitor:
                FinalSemana.selecionar_leitor_automaticamente(semanas)
            
            dirigente = FinalSemana.solicitar_dirigente_sentinela(parent)
            if dirigente is None:
                return  # usuário cancelou
            
            if not FinalSemana.solicitar_dados_usuario(semanas, parent=parent):
                return
            
            FinalSemana.criar_documento_sentinela(semanas, nome_arquivo, dirigente)
            FinalSemana.criar_documento_oradores(semanas, nome_arquivo)
            FinalSemana.salvar_historico_final_semana(semanas, nome_arquivo, dirigente)
            
        except Exception as e:
            logger.error(f"Erro ao gerar final de semana: {str(e)}")
            raise

    def buscar_titulos_meetings(url, qtd_semanas=9):
        """Busca os títulos do estudo da Sentinela na página de meetings."""
        return webscrapper.webscrapper.extrair_titulos_sentinela_meetings(url, qtd_semanas)

    def solicitar_dirigente_sentinela(parent=None):
        """
        Abre modal para selecionar o Dirigente de Sentinela entre os publicadores anciãos.
        Retorna o nome selecionado ou string vazia se confirmar em branco; None se cancelar.
        """
        resultado = [None]
        if parent is None:
            temp_root = tk.Tk()
            temp_root.withdraw()
            root_ref = temp_root
        else:
            root_ref = parent
        
        modal = ttk.Toplevel(root_ref)
        modal.title("Dirigente de Sentinela")
        modal.geometry("450x300")
        modal.transient(root_ref)
        modal.grab_set()
        
        main_container = ttk.Frame(modal, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        ttk.Label(
            main_container,
            text="Selecione o Dirigente de Sentinela",
            font=("Helvetica", 14, "bold"),
            bootstyle="primary"
        ).pack(pady=(0, 15))
        
        ancios = db_ops.listar_ancios()
        if not ancios:
            ttk.Label(main_container, text="Nenhum ancião cadastrado.", bootstyle="warning").pack(pady=10)
            combo_var = tk.StringVar(value="")
        else:
            combo_var = tk.StringVar(value=ancios[0] if ancios else "")
            combo = ttk.Combobox(
                main_container,
                textvariable=combo_var,
                values=ancios,
                width=40,
                state="readonly",
                bootstyle="primary"
            )
            combo.pack(pady=10, fill=X)
        
        def confirmar():
            resultado[0] = combo_var.get().strip() if ancios else ""
            modal.destroy()
            if parent is None:
                root_ref.destroy()
        
        def cancelar():
            resultado[0] = None
            modal.destroy()
            if parent is None:
                root_ref.destroy()
        
        btn_frame = ttk.Frame(main_container)
        btn_frame.pack(fill=X, pady=(20, 0))
        ttk.Button(btn_frame, text="Confirmar", command=confirmar, bootstyle="success", width=15).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Cancelar", command=cancelar, bootstyle="secondary", width=15).pack(side=tk.LEFT)
        
        modal.update_idletasks()
        x = (modal.winfo_screenwidth() // 2) - (450 // 2)
        y = (modal.winfo_screenheight() // 2) - (300 // 2)
        modal.geometry(f"450x300+{x}+{y}")
        modal.wait_window()
        
        return resultado[0]

    def selecionar_presidente_automaticamente(semanas):
        """
        Preenche presidente para cada semana: lista ordenada por quem fez MENOS
        participações em 'Presidente Final Semana', distribuindo um por semana (com ciclo).
        Altera semanas in-place; o modal usa a mesma referência e verá esses valores.
        """
        lista = db_ops.listar_presidentes_final_semana_ordenados_por_menos_partecipacoes()
        for i in range(len(semanas)):
            semanas[i]['presidente'] = lista[i % len(lista)] if lista else ''

    def selecionar_leitor_automaticamente(semanas):
        """
        Preenche leitor_sentinela para cada semana: lista ordenada por quem fez MENOS
        participações em 'Leitura Sentinela', distribuindo um por semana (com ciclo).
        Altera semanas in-place; o modal usa a mesma referência e verá esses valores.
        """
        lista = db_ops.listar_ordenados_por_menos_partecipacoes(
            "leitura_sentinela", "Leitura Sentinela"
        )
        for i in range(len(semanas)):
            semanas[i]['leitor_sentinela'] = lista[i % len(lista)] if lista else ''

    def _carregar_temas_discursos():
        """Carrega lista de temas de discurso de assets/temas_discursos.json. Retorna [] em erro."""
        try:
            path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'assets', 'temas_discursos.json'))
            with open(path, encoding='utf-8') as f:
                return json.load(f)
        except (FileNotFoundError, json.JSONDecodeError) as e:
            logger.warning(f"Não foi possível carregar temas_discursos.json: {e}")
            return []

    def _criar_campo_autocomplete(parent, valor_inicial, nomes_permitidos):
        """Cria Entry + Listbox com autocomplete filtrado por permissão. Retorna (container, entry)."""
        container = ttk.Frame(parent)
        entry = ttk.Entry(container, width=40, bootstyle="primary")
        entry.pack(fill=X)
        entry.insert(0, valor_inicial)
        listbox = tk.Listbox(container, height=4)
        
        def update_listbox(event):
            texto = entry.get().lower()
            listbox.delete(0, tk.END)
            if texto and nomes_permitidos:
                filtrados = [n for n in nomes_permitidos if texto in n.lower()]
                for n in filtrados[:10]:
                    listbox.insert(tk.END, n)
                if filtrados:
                    listbox.pack(fill=X, pady=(2, 0))
            else:
                listbox.pack_forget()
        
        def on_select(event):
            sel = listbox.curselection()
            if sel:
                entry.delete(0, tk.END)
                entry.insert(0, listbox.get(sel[0]))
                listbox.pack_forget()
        
        entry.bind('<KeyRelease>', update_listbox)
        listbox.bind('<<ListboxSelect>>', on_select)
        listbox.bind('<Double-Button-1>', on_select)
        return container, entry

    def solicitar_dados_usuario(semanas, parent=None):
        """
        Exibe janela para o usuário preencher Tema Discurso, Orador e Leitor (se manual).
        parent: janela pai (obrigatório para uso correto com threads - evita RuntimeError).
        Retorna True se confirmou, False se cancelou.
        IMPORTANTE: semanas é a mesma lista referenciada pela seleção automática;
        presidente e leitor_sentinela já vêm preenchidos por selecionar_*_automaticamente.
        """
        resultado = [False]
        
        # Usar parent da aplicação principal - operações GUI devem rodar na main thread
        if parent is None:
            temp_root = tk.Tk()
            temp_root.withdraw()
            root_ref = temp_root
        else:
            root_ref = parent
        
        modal = ttk.Toplevel(root_ref)
        modal.title("Preencher Dados - Reunião Final de Semana")
        modal.geometry("700x650")
        modal.transient(root_ref)
        modal.grab_set()
        
        main_container = ttk.Frame(modal, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        qtd = len(semanas)
        ttk.Label(
            main_container,
            text=f"Preencher dados das {qtd} semanas",
            font=("Helvetica", 18, "bold"),
            bootstyle="primary"
        ).pack(pady=(0, 20))
        
        scroll_frame = ttk.Frame(main_container)
        scroll_frame.pack(fill=BOTH, expand=YES)
        
        canvas = tk.Canvas(scroll_frame)
        scrollbar = ttk.Scrollbar(scroll_frame, orient="vertical", command=canvas.yview)
        scrollable = ttk.Frame(canvas)
        
        scrollable.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        temas_lista = FinalSemana._carregar_temas_discursos()
        campos = {}
        for idx, semana in enumerate(semanas):
            frame = ttk.LabelFrame(
                scrollable,
                text=f"Semana {idx + 1}: {semana.get('titulo_estudo', '')[:50]}...",
                padding=15
            )
            frame.pack(fill=X, pady=10)
            frame.grid_columnconfigure(1, weight=1)
            
            row = 0
            ttk.Label(frame, text="Tema Discurso:", bootstyle="secondary").grid(row=row, column=0, padx=(0, 10), pady=5, sticky="w")
            cont_tema, entry_tema = FinalSemana._criar_campo_autocomplete(
                frame, (semana.get('tema_discurso') or '').strip(), temas_lista
            )
            cont_tema.grid(row=row, column=1, sticky="ew", pady=5)
            row += 1
            
            ttk.Label(frame, text="Orador:", bootstyle="secondary").grid(row=row, column=0, padx=(0, 10), pady=5, sticky="w")
            entry_orador = ttk.Entry(frame, width=40, bootstyle="primary")
            entry_orador.grid(row=row, column=1, sticky="ew", pady=5)
            entry_orador.insert(0, semana.get('orador', ''))
            row += 1
            
            ttk.Label(frame, text="Presidente:", bootstyle="secondary").grid(row=row, column=0, padx=(0, 10), pady=5, sticky="w")
            cont_pres, entry_presidente = FinalSemana._criar_campo_autocomplete(
                frame, (semana.get('presidente') or '').strip(),
                db_ops.listar_ancios_e_servos()
            )
            cont_pres.grid(row=row, column=1, sticky="ew", pady=5)
            row += 1
            
            ttk.Label(frame, text="Leitor Sentinela:", bootstyle="secondary").grid(row=row, column=0, padx=(0, 10), pady=5, sticky="w")
            cont_leitor, entry_leitor = FinalSemana._criar_campo_autocomplete(
                frame, (semana.get('leitor_sentinela') or '').strip(),
                db_ops.listar_publicadores_por_permissao("leitura_sentinela")
            )
            cont_leitor.grid(row=row, column=1, sticky="ew", pady=5)
            row += 1
            
            campos[idx] = {'tema': entry_tema, 'orador': entry_orador, 'presidente': entry_presidente, 'leitor': entry_leitor}
        
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=BOTH, expand=YES)
        
        def confirmar():
            for idx, semana in enumerate(semanas):
                c = campos.get(idx, {})
                semana['tema_discurso'] = c.get('tema', ttk.Entry()).get().strip()
                semana['orador'] = c.get('orador', ttk.Entry()).get().strip()
                semana['presidente'] = c.get('presidente', ttk.Entry()).get().strip()
                semana['leitor_sentinela'] = c.get('leitor', ttk.Entry()).get().strip()
            resultado[0] = True
            modal.destroy()
            if parent is None:
                root_ref.destroy()
        
        def cancelar():
            resultado[0] = False
            modal.destroy()
            if parent is None:
                root_ref.destroy()
        
        btn_frame = ttk.Frame(main_container)
        btn_frame.pack(fill=X, pady=(20, 0))
        ttk.Button(btn_frame, text="Gerar Documento", command=confirmar, bootstyle="success", width=20).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Cancelar", command=cancelar, bootstyle="secondary", width=20).pack(side=tk.LEFT)
        
        modal.update_idletasks()
        x = (modal.winfo_screenwidth() // 2) - (700 // 2)
        y = (modal.winfo_screenheight() // 2) - (650 // 2)
        modal.geometry(f"700x650+{x}+{y}")
        modal.wait_window()
        
        return resultado[0]

    def _aplicar_placeholders(doc, placeholders):
        """Aplica dicionário de placeholders em parágrafos e tabelas do documento."""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for ph, valor in placeholders.items():
                    if ph in run.text:
                        run.text = run.text.replace(ph, str(valor))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for ph, valor in placeholders.items():
                                if ph in run.text:
                                    run.text = run.text.replace(ph, str(valor))

    def criar_documento_sentinela(semanas, nome_arquivo, dirigente):
        """
        Gera o documento Sentinela (w_dirigente, w1..w9 leitor e titulo).
        O template é apenas LIDO de Templates/; a saída é salva em documentosCriados/.
        """
        template_path = "Templates/template_final_semana_sentinela.docx"
        try:
            doc = Document(template_path)
        except Exception as e:
            raise FileNotFoundError(f"Template não encontrado: {template_path}. Crie o arquivo com w_dirigente, w1_titulo, w1_leitor até w9.")
        
        placeholders = {'w_dirigente': dirigente}
        for i, semana in enumerate(semanas):
            n = i + 1
            placeholders[f'w{n}_titulo'] = semana.get('titulo_estudo', '')
            placeholders[f'w{n}_leitor'] = semana.get('leitor_sentinela', '')
        
        FinalSemana._aplicar_placeholders(doc, placeholders)
        nome_base = nome_arquivo.replace('.docx', '').strip() if nome_arquivo else "documento"
        nome_sentinela = nome_base + "_sentinela.docx"
        path = "documentosCriados/" + nome_sentinela
        doc.save(path)
        logger.info(f"Documento Sentinela gerado: {path}")
        subprocess.run(["start", path], shell=True)

    def criar_documento_oradores(semanas, nome_arquivo):
        """
        Gera o documento Oradores (w1..w9 presidente, orador e tema).
        O template é apenas LIDO de Templates/; a saída é salva em documentosCriados/.
        """
        template_path = "Templates/template_final_semana_oradores.docx"
        try:
            doc = Document(template_path)
        except Exception as e:
            raise FileNotFoundError(f"Template não encontrado: {template_path}. Crie o arquivo com w1_presidente, w1_orador, w1_tema até w9.")
        
        placeholders = {}
        for i, semana in enumerate(semanas):
            n = i + 1
            placeholders[f'w{n}_presidente'] = semana.get('presidente', '')
            placeholders[f'w{n}_orador'] = semana.get('orador', '')
            placeholders[f'w{n}_tema'] = semana.get('tema_discurso', '')
        
        FinalSemana._aplicar_placeholders(doc, placeholders)
        nome_base = nome_arquivo.replace('.docx', '').strip() if nome_arquivo else "documento"
        nome_oradores = nome_base + "_oradores.docx"
        path = "documentosCriados/" + nome_oradores
        doc.save(path)
        logger.info(f"Documento Oradores gerado: {path}")
        subprocess.run(["start", path], shell=True)

    def salvar_historico_final_semana(semanas, nome_arquivo, dirigente=""):
        """Salva o histórico na collection reunioes_final_semana."""
        now = datetime.datetime.now()
        dados = {
            'ano': now.year,
            'mes': now.month,
            'data_criacao': now.isoformat(),
            'nome_arquivo': nome_arquivo,
            'dirigente': dirigente,
            'semanas': semanas
        }
        resultado = db_ops.salvar_reuniao_final_semana(dados)
        if not resultado.get('success'):
            logger.warning(f"Erro ao salvar histórico: {resultado.get('message', '')}")
