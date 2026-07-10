"""designacoes_salao

Tela extraída de layout.py (SDD 03) — mixin de ModernApp.
Código movido verbatim; `self` e navegação inalterados.
"""
from views._shared import *  # noqa: F401,F403


class DesignacoesMixin:
    def designacoes_salao(self):
        from process.designacoes_salao import DesignacoesSalao

        win = ttk.Toplevel(self.root)
        win.title("Designações Salão")
        win.geometry("1150x720")

        main = ttk.Frame(win, padding=20)
        main.pack(fill=BOTH, expand=YES)

        ttk.Label(main, text="Designações Salão",
                  font=("Helvetica", 20, "bold"), bootstyle="primary").pack(anchor=W, pady=(0, 15))

        # ── Configurações ────────────────────────────────────────────────
        cfg = ttk.LabelFrame(main, text="Configurações", padding=10)
        cfg.pack(fill=X, pady=(0, 10))

        mapa_dia = {"(nenhum)": None, "Segunda": 0, "Terça": 1, "Quarta": 2, "Quinta": 3, "Sexta": 4}
        mapa_fds = {"(nenhum)": None, "Sábado": 5, "Domingo": 6}

        ttk.Label(cfg, text="Dia Meio de Semana:").grid(row=0, column=0, padx=(0, 5), pady=5, sticky=W)
        dia_semana_var = ttk.StringVar(value="(nenhum)")
        ttk.Combobox(cfg, textvariable=dia_semana_var,
                     values=list(mapa_dia.keys()), width=12, state="readonly",
                     bootstyle="primary").grid(row=0, column=1, padx=(0, 20), pady=5, sticky=W)

        ttk.Label(cfg, text="Dia Final de Semana:").grid(row=0, column=2, padx=(0, 5), pady=5, sticky=W)
        dia_fds_var = ttk.StringVar(value="(nenhum)")
        ttk.Combobox(cfg, textvariable=dia_fds_var,
                     values=list(mapa_fds.keys()), width=12, state="readonly",
                     bootstyle="primary").grid(row=0, column=3, padx=(0, 20), pady=5, sticky=W)

        ttk.Button(cfg, text="Gerar Automaticamente", bootstyle="success",
                   command=lambda: gerar_auto()).grid(row=0, column=4, padx=5, pady=5)
        ttk.Button(cfg, text="Carregar Dados Salvos", bootstyle="info",
                   command=lambda: carregar_salvos()).grid(row=0, column=5, padx=5, pady=5)

        # ── Tabela ────────────────────────────────────────────────────────
        tabela_frame = ttk.Frame(main)
        tabela_frame.pack(fill=BOTH, expand=YES, pady=(0, 10))

        cols = ("data", "tipo", "audio", "video", "microfone", "indicadores")
        tabela = ttk.Treeview(tabela_frame, columns=cols, show="tree headings",
                              bootstyle="primary", height=18,
                              displaycolumns=("data", "audio", "video", "microfone", "indicadores"))
        tabela.heading("#0",          text="Mês")
        tabela.heading("data",        text="Data")
        tabela.heading("audio",       text="Áudio")
        tabela.heading("video",       text="Vídeo")
        tabela.heading("microfone",   text="Microfone")
        tabela.heading("indicadores", text="Indicadores")
        tabela.column("#0",          width=140, anchor=W)
        tabela.column("data",        width=80,  anchor=CENTER)
        tabela.column("audio",       width=150, anchor=W)
        tabela.column("video",       width=150, anchor=W)
        tabela.column("microfone",   width=210, anchor=W)
        tabela.column("indicadores", width=210, anchor=W)

        sb_v = ttk.Scrollbar(tabela_frame, orient="vertical",
                             command=tabela.yview, bootstyle="primary-round")
        sb_v.pack(side=RIGHT, fill=Y)
        tabela.configure(yscrollcommand=sb_v.set)
        sb_h = ttk.Scrollbar(tabela_frame, orient="horizontal",
                             command=tabela.xview, bootstyle="primary-round")
        sb_h.pack(side=BOTTOM, fill=X)
        tabela.configure(xscrollcommand=sb_h.set)
        tabela.pack(fill=BOTH, expand=YES)

        # ── Botões ────────────────────────────────────────────────────────
        btn_frame = ttk.Frame(main)
        btn_frame.pack(fill=X)

        ttk.Button(btn_frame, text="Salvar", bootstyle="primary",
                   command=lambda: salvar()).pack(side=LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Exportar DOCX", bootstyle="success",
                   command=lambda: exportar_docx()).pack(side=LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Excluir Mês Selecionado", bootstyle="danger",
                   command=lambda: excluir_mes()).pack(side=LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Fechar", bootstyle="secondary",
                   command=win.destroy).pack(side=RIGHT)

        # ── Callbacks ─────────────────────────────────────────────────────
        def _popula_tabela(designacoes):
            import datetime as _dt
            from collections import defaultdict
            for row in tabela.get_children():
                tabela.delete(row)
            MESES_PT = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                        6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                        10:"Outubro",11:"Novembro",12:"Dezembro"}
            por_mes = defaultdict(list)
            for d in designacoes:
                try:
                    dt = _dt.datetime.strptime(d["data"], "%d/%m/%Y")
                    por_mes[(dt.year, dt.month)].append(d)
                except ValueError:
                    por_mes[(0, 0)].append(d)
            for (ano, mes) in sorted(por_mes.keys()):
                mes_label = f"{MESES_PT.get(mes, str(mes))} {ano}" if mes else "?"
                mes_iid = f"mes_{ano}_{mes}"
                tabela.insert("", "end", iid=mes_iid, text=mes_label, open=True, tags=("mes_header",))
                for d in por_mes[(ano, mes)]:
                    tabela.insert(mes_iid, "end", values=(
                        d["data"], d.get("tipo", ""),
                        d.get("audio", ""), d.get("video", ""),
                        d.get("microfone", ""), d.get("indicadores", "")
                    ))
            tabela.tag_configure("mes_header", font=("Helvetica", 10, "bold"), foreground="#1a56a0")

        def carregar_salvos():
            from ttkbootstrap.dialogs import Messagebox
            docs = designacao_service.listar()
            if not docs:
                Messagebox.show_info("Nenhuma designação salva encontrada.", "Carregar")
                return
            todas = []
            for doc in docs:
                todas.extend(doc.get("semanas", []))
            if not todas:
                Messagebox.show_info("Nenhuma designação salva encontrada.", "Carregar")
                return
            todas.sort(key=lambda x: __import__('datetime').datetime.strptime(x["data"], "%d/%m/%Y"))
            _popula_tabela(todas)

        def gerar_auto():
            import datetime as _dt
            from ttkbootstrap.dialogs import Messagebox
            idx_dia = mapa_dia.get(dia_semana_var.get())
            idx_fds = mapa_fds.get(dia_fds_var.get())
            if idx_dia is None and idx_fds is None:
                Messagebox.show_warning("Selecione ao menos um dia de reunião.", "Atenção")
                return
            MESES_PT = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                        6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                        10:"Outubro",11:"Novembro",12:"Dezembro"}
            dlg = ttk.Toplevel(win)
            dlg.title("Selecionar Meses")
            dlg.grab_set()
            dlg.resizable(False, False)
            ttk.Label(dlg, text="Selecione os meses para gerar:",
                      font=("Helvetica", 11, "bold"), bootstyle="primary").pack(padx=20, pady=(15, 8))
            hoje = _dt.date.today()
            opcoes = []
            for i in range(12):
                m = (hoje.month - 1 + i) % 12 + 1
                a = hoje.year + (hoje.month - 1 + i) // 12
                opcoes.append((a, m))
            chk_vars = {}
            chk_frame = ttk.Frame(dlg, padding=5)
            chk_frame.pack(padx=20, pady=5)
            for i, (a, m) in enumerate(opcoes):
                var = ttk.BooleanVar(value=False)
                chk_vars[(a, m)] = var
                ttk.Checkbutton(chk_frame, text=f"{MESES_PT[m]} {a}",
                                variable=var, bootstyle="primary").grid(
                    row=i // 3, column=i % 3, sticky=W, padx=12, pady=3)
            def confirmar_gerar():
                selecionados = [(a, m) for (a, m), v in chk_vars.items() if v.get()]
                if not selecionados:
                    Messagebox.show_warning("Selecione ao menos um mês.", "Atenção", parent=dlg)
                    return
                dlg.destroy()
                selecionados.sort()
                datas = DesignacoesSalao.gerar_datas_meses(idx_dia, idx_fds, selecionados)
                if not datas:
                    Messagebox.show_warning("Nenhuma data gerada.", "Atenção", parent=win)
                    return
                _abrir_impedimentos(selecionados, datas)

            def _gerar_e_popular(datas, impedimentos):
                designacoes = DesignacoesSalao.gerar_designacoes(datas, impedimentos)
                _popula_tabela(designacoes)

            def _abrir_impedimentos(selecionados, datas):
                import calendar as _cal

                # União dos irmãos com permissão de salão (dedup, preserva ordem)
                irmaos, vistos = [], set()
                for papel in ("audio_video", "microfone", "indicador"):
                    for n in designacao_service.listar_candidatos(papel):
                        if n and n not in vistos:
                            vistos.add(n)
                            irmaos.append(n)

                if not irmaos:
                    _gerar_e_popular(datas, {})
                    return

                datas_reuniao = {d["data"] for d in datas}
                impedimentos = {}          # {nome: set("DD/MM/AAAA")}
                estado = {"nome": None}    # irmão atualmente selecionado

                imp = ttk.Toplevel(win)
                imp.title("Impedimentos dos Irmãos")
                imp.grab_set()
                imp.geometry("820x600")

                ttk.Label(imp, text="Algum irmão tem impedimento em alguns dias?",
                          font=("Helvetica", 13, "bold"), bootstyle="primary").pack(
                    anchor=W, padx=20, pady=(15, 2))
                ttk.Label(imp, text="Selecione o irmão à esquerda e marque no calendário "
                          "os dias em que ele NÃO pode trabalhar.",
                          bootstyle="secondary").pack(anchor=W, padx=20, pady=(0, 10))

                corpo = ttk.Frame(imp, padding=(20, 0))
                corpo.pack(fill=BOTH, expand=YES)

                # ── Esquerda: lista de irmãos ──────────────────────────────
                esq = ttk.LabelFrame(corpo, text="Irmãos", padding=5)
                esq.pack(side=LEFT, fill=Y, padx=(0, 12))
                lst = ttk.Treeview(esq, show="tree", height=22, bootstyle="primary")
                lst.column("#0", width=210, anchor=W)
                sb_l = ttk.Scrollbar(esq, orient="vertical", command=lst.yview,
                                     bootstyle="primary-round")
                sb_l.pack(side=RIGHT, fill=Y)
                lst.configure(yscrollcommand=sb_l.set)
                lst.pack(side=LEFT, fill=Y)
                for n in irmaos:
                    lst.insert("", "end", iid=n, text=n)

                # ── Direita: calendário rolável ────────────────────────────
                dir_ = ttk.LabelFrame(corpo, text="Calendário", padding=5)
                dir_.pack(side=LEFT, fill=BOTH, expand=YES)
                canvas = tk.Canvas(dir_, highlightthickness=0)
                sb_c = ttk.Scrollbar(dir_, orient="vertical", command=canvas.yview,
                                     bootstyle="primary-round")
                cal_frame = ttk.Frame(canvas)
                cal_frame.bind("<Configure>",
                               lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
                canvas.create_window((0, 0), window=cal_frame, anchor="nw")
                canvas.configure(yscrollcommand=sb_c.set)
                sb_c.pack(side=RIGHT, fill=Y)
                canvas.pack(side=LEFT, fill=BOTH, expand=YES)

                MESES_LOCAL = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                               6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                               10:"Outubro",11:"Novembro",12:"Dezembro"}
                dias_botoes = []   # [(btn, "DD/MM/AAAA", is_reuniao)]

                def _estilo_dia(data_str, is_reuniao, marcado):
                    if marcado:
                        return "danger"
                    return "info-outline" if is_reuniao else "secondary-outline"

                def _repintar():
                    nome = estado["nome"]
                    marcados = impedimentos.get(nome, set())
                    for btn, data_str, is_reuniao in dias_botoes:
                        st = "disabled" if nome is None else "normal"
                        btn.configure(bootstyle=_estilo_dia(data_str, is_reuniao,
                                                            data_str in marcados),
                                      state=st)

                def _toggle(data_str):
                    nome = estado["nome"]
                    if nome is None:
                        return
                    s = impedimentos.setdefault(nome, set())
                    if data_str in s:
                        s.discard(data_str)
                    else:
                        s.add(data_str)
                    if not s:
                        impedimentos.pop(nome, None)
                    lst.item(nome, text=f"{nome} ({len(s)})" if s else nome)
                    _repintar()

                cal = _cal.Calendar(firstweekday=6)  # domingo primeiro
                header = ["D", "S", "T", "Q", "Q", "S", "S"]
                for (ano, mes) in selecionados:
                    mf = ttk.LabelFrame(cal_frame,
                                        text=f"{MESES_LOCAL[mes]} {ano}", padding=6)
                    mf.pack(fill=X, pady=6, padx=2)
                    for c, h in enumerate(header):
                        ttk.Label(mf, text=h, width=4, anchor=CENTER,
                                  bootstyle="secondary").grid(row=0, column=c, padx=1, pady=1)
                    for r, semana in enumerate(cal.monthdayscalendar(ano, mes), start=1):
                        for c, dia in enumerate(semana):
                            if dia == 0:
                                continue
                            data_str = f"{dia:02d}/{mes:02d}/{ano}"
                            is_reuniao = data_str in datas_reuniao
                            b = ttk.Button(mf, text=str(dia), width=4,
                                           bootstyle=_estilo_dia(data_str, is_reuniao, False),
                                           command=lambda d=data_str: _toggle(d))
                            b.grid(row=r, column=c, padx=1, pady=1)
                            b.configure(state="disabled")
                            dias_botoes.append((b, data_str, is_reuniao))

                def _on_sel(_evt=None):
                    sel = lst.selection()
                    estado["nome"] = sel[0] if sel else None
                    _repintar()
                lst.bind("<<TreeviewSelect>>", _on_sel)

                # ── Botões ─────────────────────────────────────────────────
                bf = ttk.Frame(imp)
                bf.pack(fill=X, padx=20, pady=12)

                def _confirmar():
                    imp.destroy()
                    _gerar_e_popular(datas, impedimentos)

                def _pular():
                    imp.destroy()
                    _gerar_e_popular(datas, {})

                ttk.Button(bf, text="Gerar", bootstyle="success",
                           command=_confirmar).pack(side=LEFT, padx=(0, 10))
                ttk.Button(bf, text="Gerar sem impedimentos", bootstyle="secondary",
                           command=_pular).pack(side=LEFT, padx=(0, 10))
                ttk.Button(bf, text="Cancelar", bootstyle="danger-outline",
                           command=imp.destroy).pack(side=RIGHT)

                imp.update_idletasks()
                x = win.winfo_x() + (win.winfo_width() - imp.winfo_width()) // 2
                y = win.winfo_y() + (win.winfo_height() - imp.winfo_height()) // 2
                imp.geometry(f"+{max(x, 0)}+{max(y, 0)}")
            btn_f = ttk.Frame(dlg)
            btn_f.pack(padx=20, pady=15)
            ttk.Button(btn_f, text="Gerar", bootstyle="success",
                       command=confirmar_gerar).pack(side=LEFT, padx=(0, 10))
            ttk.Button(btn_f, text="Cancelar", bootstyle="secondary",
                       command=dlg.destroy).pack(side=LEFT)
            dlg.update_idletasks()
            x = win.winfo_x() + (win.winfo_width() - dlg.winfo_reqwidth()) // 2
            y = win.winfo_y() + (win.winfo_height() - dlg.winfo_reqheight()) // 2
            dlg.geometry(f"+{x}+{y}")

        def salvar():
            import datetime as _dt
            from collections import defaultdict
            from ttkbootstrap.dialogs import Messagebox

            # Coletar dados da tabela
            por_mes = defaultdict(list)
            todas_linhas = []
            for mes_iid in tabela.get_children():
                for iid in tabela.get_children(mes_iid):
                    vals = tabela.item(iid)["values"]
                    data_str, tipo, audio, video, mic, ind = (str(v) for v in vals)
                    try:
                        d = _dt.datetime.strptime(data_str, "%d/%m/%Y")
                    except ValueError:
                        continue
                    row = {"data": data_str, "tipo": tipo, "audio": audio,
                           "video": video, "microfone": mic, "indicadores": ind}
                    por_mes[(d.year, d.month)].append(row)
                    todas_linhas.append(row)

            if not por_mes:
                Messagebox.show_warning("Nenhum dado para salvar.", "Atenção")
                return

            # ── Janela de pré-visualização ────────────────────────────────
            prev = ttk.Toplevel(win)
            prev.title("Validar Designações")
            prev.grab_set()
            prev_w, prev_h = 1100, 600
            px = (prev.winfo_screenwidth() // 2) - (prev_w // 2)
            py = (prev.winfo_screenheight() // 2) - (prev_h // 2)
            prev.geometry(f"{prev_w}x{prev_h}+{px}+{py}")

            ttk.Label(prev, text="Revise as designações antes de salvar:",
                      font=("Helvetica", 13, "bold"),
                      bootstyle="primary").pack(anchor=W, padx=20, pady=(15, 5))

            meses_pt = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                        6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                        10:"Outubro",11:"Novembro",12:"Dezembro"}
            info_meses = ", ".join(
                f"{meses_pt.get(m, m)}/{a}" for a, m in sorted(por_mes.keys())
            )
            ttk.Label(prev,
                      text=f"{len(todas_linhas)} datas  |  Meses: {info_meses}",
                      bootstyle="secondary").pack(anchor=W, padx=20, pady=(0, 8))

            # Treeview de prévia
            prev_frame = ttk.Frame(prev)
            prev_frame.pack(fill=BOTH, expand=YES, padx=20)

            prev_cols = ("data", "audio", "video", "microfone", "indicadores")
            prev_tree = ttk.Treeview(prev_frame, columns=prev_cols,
                                     show="headings", bootstyle="primary", height=16)
            prev_tree.heading("data",        text="Data")
            prev_tree.heading("audio",       text="Áudio")
            prev_tree.heading("video",       text="Vídeo")
            prev_tree.heading("microfone",   text="Microfone")
            prev_tree.heading("indicadores", text="Indicadores")
            prev_tree.column("data",        width=85,  anchor=CENTER)
            prev_tree.column("audio",       width=155, anchor=W)
            prev_tree.column("video",       width=155, anchor=W)
            prev_tree.column("microfone",   width=215, anchor=W)
            prev_tree.column("indicadores", width=215, anchor=W)

            prev_sb = ttk.Scrollbar(prev_frame, orient="vertical",
                                    command=prev_tree.yview, bootstyle="primary-round")
            prev_sb.pack(side=RIGHT, fill=Y)
            prev_tree.configure(yscrollcommand=prev_sb.set)
            prev_tree.pack(fill=BOTH, expand=YES)

            for row in todas_linhas:
                prev_tree.insert("", "end", values=(
                    row["data"], row["audio"], row["video"],
                    row["microfone"], row["indicadores"]
                ))

            # Botões confirmar / cancelar
            btn_prev = ttk.Frame(prev)
            btn_prev.pack(fill=X, padx=20, pady=15)

            def confirmar_salvar():
                prev.destroy()
                for (ano, mes), semanas in por_mes.items():
                    dados = {
                        "ano": ano, "mes": mes,
                        "dia_semana": dia_semana_var.get(),
                        "dia_fds": dia_fds_var.get(),
                        "semanas": semanas
                    }
                    res = designacao_service.salvar(dados)
                    if not res.get("success"):
                        Messagebox.show_error(res.get("message", "Erro ao salvar."), "Erro", parent=win)
                        return
                Messagebox.show_info("Designações salvas com sucesso!", "Sucesso", parent=win)

            ttk.Button(btn_prev, text="Confirmar e Salvar", bootstyle="success",
                       command=confirmar_salvar).pack(side=LEFT, padx=(0, 10))
            ttk.Button(btn_prev, text="Cancelar", bootstyle="secondary",
                       command=prev.destroy).pack(side=LEFT)

        def excluir_mes():
            from ttkbootstrap.dialogs import Messagebox
            MESES_PT = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                        6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                        10:"Outubro",11:"Novembro",12:"Dezembro"}
            sel = tabela.selection()
            if not sel:
                Messagebox.show_warning("Selecione um mês ou data.", "Atenção", parent=win)
                return
            iid = sel[0]
            parent = tabela.parent(iid)
            mes_iid = iid if parent == "" else parent
            try:
                _, ano_str, mes_str = mes_iid.split("_")
                ano, mes = int(ano_str), int(mes_str)
            except (ValueError, AttributeError):
                Messagebox.show_warning("Não foi possível identificar o mês.", "Atenção", parent=win)
                return
            confirm = Messagebox.yesno(
                f"Excluir todas as designações de {MESES_PT.get(mes, mes)}/{ano}?\nEsta ação remove o histórico dos irmãos.",
                "Confirmar Exclusão", parent=win
            )
            if confirm not in (None, "No", "Não", "Cancelar"):
                res = designacao_service.excluir(ano, mes)
                tabela.delete(mes_iid)
                msg = res.get("message", "")
                if res.get("success"):
                    Messagebox.show_info(msg or "Excluído.", "Sucesso", parent=win)
                elif msg and msg not in ("Designações não encontradas",):
                    Messagebox.show_error(msg, "Erro", parent=win)

        def exportar_docx():
            import datetime as _dt
            import os
            from tkinter import filedialog
            from ttkbootstrap.dialogs import Messagebox
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            mes_iids = tabela.get_children()
            if not mes_iids:
                Messagebox.show_warning("Nenhum dado na tabela.", "Atenção")
                return

            MESES_PT = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                        6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                        10:"Outubro",11:"Novembro",12:"Dezembro"}

            mes_info = []
            for iid in mes_iids:
                try:
                    _, ano_str, mes_str = iid.split("_")
                    mes_info.append((int(ano_str), int(mes_str), iid))
                except (ValueError, AttributeError):
                    continue

            if not mes_info:
                Messagebox.show_warning("Nenhum dado na tabela.", "Atenção")
                return

            dlg = ttk.Toplevel(win)
            dlg.title("Exportar DOCX — Selecionar Meses")
            dlg.grab_set()
            dlg.resizable(False, False)
            ttk.Label(dlg, text="Selecione os meses para exportar:",
                      font=("Helvetica", 11, "bold"), bootstyle="primary").pack(padx=20, pady=(15, 8))
            chk_vars = {}
            chk_frame = ttk.Frame(dlg, padding=5)
            chk_frame.pack(padx=20, pady=5)
            for i, (ano, mes, iid) in enumerate(mes_info):
                var = ttk.BooleanVar(value=True)
                chk_vars[iid] = (ano, mes, var)
                ttk.Checkbutton(chk_frame, text=f"{MESES_PT[mes]} {ano}",
                                variable=var, bootstyle="primary").grid(
                    row=i // 3, column=i % 3, sticky=W, padx=12, pady=3)

            def fazer_exportar():
                rows = []
                for iid, (ano, mes, var) in chk_vars.items():
                    if not var.get():
                        continue
                    for child_iid in tabela.get_children(iid):
                        vals = tabela.item(child_iid)["values"]
                        rows.append({
                            "data":        str(vals[0]),
                            "audio":       str(vals[2]),
                            "video":       str(vals[3]),
                            "microfone":   str(vals[4]),
                            "indicadores": str(vals[5]),
                        })
                if not rows:
                    Messagebox.show_warning("Nenhum mês selecionado.", "Atenção", parent=dlg)
                    return
                dlg.destroy()

                caminho = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Document", "*.docx")],
                    initialfile="Designacoes_Salao.docx",
                    title="Salvar DOCX"
                )
                if not caminho:
                    return

                doc = Document()

                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = Cm(29.7), Cm(21.0)
                section.top_margin    = Cm(1.2)
                section.bottom_margin = Cm(1.2)
                section.left_margin   = Cm(1.5)
                section.right_margin  = Cm(1.5)

                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_title = p_title.add_run("Designações Salão")
                run_title.font.size = Pt(13)
                run_title.font.bold = True

                col_widths = [Cm(2.2), Cm(4.8), Cm(4.8), Cm(6.5), Cm(6.5)]
                headers    = ["Data", "Áudio", "Vídeo", "Microfone", "Indicadores"]

                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"

                hdr = table.rows[0].cells
                for i, (txt, w) in enumerate(zip(headers, col_widths)):
                    hdr[i].width = w
                    p = hdr[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(txt)
                    run.font.size = Pt(9)
                    run.font.bold = True
                    tc = hdr[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'BDD7EE')
                    tcPr.append(shd)

                meses_pt = {1:"JANEIRO",2:"FEVEREIRO",3:"MARÇO",4:"ABRIL",5:"MAIO",
                            6:"JUNHO",7:"JULHO",8:"AGOSTO",9:"SETEMBRO",
                            10:"OUTUBRO",11:"NOVEMBRO",12:"DEZEMBRO"}

                def _set_cell_fill(cell, hex_color):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), hex_color)
                    tcPr.append(shd)

                fill_data = ['FFFFFF', 'EBF3FB']
                prev_month = None
                month_color_idx = 0

                for row_data in rows:
                    try:
                        d = _dt.datetime.strptime(row_data["data"], "%d/%m/%Y")
                        cur_month = (d.year, d.month)
                    except ValueError:
                        cur_month = None

                    if cur_month != prev_month:
                        month_color_idx = 1 - month_color_idx
                        prev_month = cur_month

                        if cur_month:
                            mes_label = f"{meses_pt.get(cur_month[1], '')}  —  {cur_month[0]}"
                            hdr_row = table.add_row().cells
                            merged = hdr_row[0].merge(hdr_row[1]).merge(hdr_row[2]).merge(hdr_row[3]).merge(hdr_row[4])
                            p_hdr = merged.paragraphs[0]
                            p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_hdr = p_hdr.add_run(mes_label)
                            run_hdr.font.size = Pt(9)
                            run_hdr.font.bold = True
                            run_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            _set_cell_fill(merged, '2F5496')

                    row_cells = table.add_row().cells
                    values_row = [
                        row_data["data"], row_data["audio"], row_data["video"],
                        row_data["microfone"], row_data["indicadores"]
                    ]
                    for i, (val, w) in enumerate(zip(values_row, col_widths)):
                        row_cells[i].width = w
                        p = row_cells[i].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(val)
                        run.font.size = Pt(8)
                        fill = fill_data[month_color_idx]
                        if fill != 'FFFFFF':
                            _set_cell_fill(row_cells[i], fill)

                doc.save(caminho)
                if Messagebox.yesno(f"DOCX salvo em:\n{caminho}\n\nAbrir arquivo?", "Exportado") == "Yes":
                    os.startfile(caminho)

            btn_f = ttk.Frame(dlg)
            btn_f.pack(padx=20, pady=15)
            ttk.Button(btn_f, text="Exportar", bootstyle="success",
                       command=fazer_exportar).pack(side=LEFT, padx=(0, 10))
            ttk.Button(btn_f, text="Cancelar", bootstyle="secondary",
                       command=dlg.destroy).pack(side=LEFT)
            dlg.update_idletasks()
            x = win.winfo_x() + (win.winfo_width() - dlg.winfo_reqwidth()) // 2
            y = win.winfo_y() + (win.winfo_height() - dlg.winfo_reqheight()) // 2
            dlg.geometry(f"+{x}+{y}")

        def on_double_click(event):
            item = tabela.identify_row(event.y)
            col  = tabela.identify_column(event.x)
            if not item:
                return
            # Ignorar nós de mês e coluna de árvore/data
            if tabela.parent(item) == "" or col in ("#0", "#1"):
                return
            # #2=audio, #3=video, #4=microfone, #5=indicadores
            col_to_val = {"#2": 2, "#3": 3, "#4": 4, "#5": 5}
            col_labels  = {"#2": "Áudio", "#3": "Vídeo", "#4": "Microfone", "#5": "Indicadores"}
            if col not in col_to_val:
                return
            values_idx = col_to_val[col]
            current = tabela.item(item)["values"]
            popup = ttk.Toplevel(win)
            popup.title(f"Editar — {col_labels.get(col, '')}")
            popup.resizable(False, False)
            popup.grab_set()
            entry_var = ttk.StringVar(value=str(current[values_idx]))
            ttk.Entry(popup, textvariable=entry_var, width=40,
                      bootstyle="primary").pack(padx=15, pady=(15, 5))
            if col in ("#4", "#5"):
                ttk.Label(popup, text='Formato: "Nome1 / Nome2"',
                          bootstyle="secondary", font=("Helvetica", 9)).pack()
            def confirmar(event=None):
                vals = list(tabela.item(item)["values"])
                vals[values_idx] = entry_var.get()
                tabela.item(item, values=vals)
                popup.destroy()
            ttk.Button(popup, text="OK", command=confirmar,
                       bootstyle="primary").pack(pady=(8, 15))
            popup.bind("<Return>", confirmar)

        tabela.bind("<Double-1>", on_double_click)

        # Carregar dados salvos automaticamente ao abrir
        carregar_salvos()

        # Centralizar janela
        win.update_idletasks()
        x = (win.winfo_screenwidth() // 2) - (1150 // 2)
        y = (win.winfo_screenheight() // 2) - (720 // 2)
        win.geometry(f"1150x720+{x}+{y}")

