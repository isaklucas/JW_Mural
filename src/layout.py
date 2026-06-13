import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent))
if '--init-db' in sys.argv:
    from database.init_db import init_mongodb
    _ok = init_mongodb()
    sys.exit(0 if _ok else 1)

import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import tkinter as tk
import webbrowser
import time
import process.s140 as s140
import process.final_semana as final_semana
from PIL import Image, ImageTk
import os
from database import post, getAllPub, delete, listar_reunioes_final_semana, buscar_reuniao_final_semana, excluir_reuniao_final_semana, db_ops
from database.db_operations import DatabaseOperations
import util.janelas as janelas
from util.startup_manager import initialize_application
import sys
import logging
import datetime
import threading

logger = logging.getLogger(__name__)

class ModernApp:
    def __init__(self, root):
        self.root = root
        # Configuração da janela principal com tema moderno
        self.root.iconbitmap("assets/icon.ico") if os.path.exists("assets/icon.ico") else None
        
        # Configuração do container principal
        self.main_frame = ttk.Frame(self.root, padding=20)
        self.main_frame.pack(fill=BOTH, expand=YES)
        
        # Título principal
        title_frame = ttk.Frame(self.main_frame)
        title_frame.pack(fill=X, pady=(0, 20))
        
        title_label = ttk.Label(
            title_frame,
            text="JW Mural",
            font=("Helvetica", 24, "bold"),
            bootstyle="primary"
        )
        title_label.pack()
        
        subtitle_label = ttk.Label(
            title_frame,
            text="Sistema de Gerenciamento de Reuniões",
            font=("Helvetica", 12),
            bootstyle="secondary"
        )
        subtitle_label.pack()
        
        # Container para os cards
        self.cards_frame = ttk.Frame(self.main_frame, bootstyle="light")
        self.cards_frame.pack(fill=BOTH, expand=YES)
        
        # Criação dos cards
        # Linha 0: Criar reuniões
        self.create_card(
            "Criar Reunião de Meio de Semana",
            "Gere documentos para reuniões de meio de semana",
            "primary",
            self.criar_quadro_de_anuncio,
            0, 0
        )
        self.create_card(
            "Criar Reunião Final de Semana",
            "Gere documentos para reuniões de final de semana",
            "primary",
            self.criar_reuniao_final_semana,
            0, 1
        )
        
        # Linha 1: Publicadores
        self.create_card(
            "Publicadores",
            "Gerencie a lista de publicadores",
            "info",
            self.publicadores,
            1, 0
        )
        self.create_card(
            "Histórico de Publicadores",
            "Visualize o histórico de partes dos publicadores",
            "success",
            self.historico_publicadores,
            1, 1
        )
        
        # Linha 2: Históricos de reuniões
        self.create_card(
            "Histórico Reunião Meio de Semana",
            "Visualize o histórico de reuniões de meio de semana",
            "warning",
            self.historico,
            2, 0
        )
        self.create_card(
            "Histórico Reunião Final de Semana",
            "Visualize o histórico de reuniões de final de semana",
            "secondary",
            self.historico_final_semana,
            2, 1
        )
        
        # Linha 3: Designações Salão + Dashboards
        self.create_card(
            "Designações Salão",
            "Gerencie áudio, vídeo, microfone e indicadores",
            "info",
            self.designacoes_salao,
            3, 0
        )
        self.create_card(
            "Dashboards",
            "Visualize estatísticas e gráficos",
            "danger",
            self.dashboards,
            3, 1
        )
        
        # Rodapé
        footer_frame = ttk.Frame(self.main_frame)
        footer_frame.pack(fill=X, pady=(20, 0))
        
        version_label = ttk.Label(
            footer_frame,
            text="Versão 1.0",
            bootstyle="secondary"
        )
        version_label.pack(side=RIGHT)

        gear_btn = ttk.Button(
            self.main_frame,
            text="⚙",
            bootstyle="secondary-outline",
            width=3,
            command=self.abrir_configuracoes
        )
        gear_btn.place(relx=1.0, rely=0.0, anchor="ne")

    def create_card(self, title, description, style, command, row, col, columnspan=1):
        """Cria um card estilizado com título, descrição e botão"""
        # Card principal com preenchimento maior e bordas arredondadas
        card = ttk.Frame(
            self.cards_frame,
            bootstyle="light",  # Alterado para light (branco)
            padding=20
        )
        card.grid(row=row, column=col, columnspan=columnspan, padx=10, pady=10, sticky="nsew")
        
        # Configurar o grid para expandir igualmente
        for c in range(col, col + columnspan):
            self.cards_frame.grid_columnconfigure(c, weight=1)
        self.cards_frame.grid_rowconfigure(row, weight=1)
        
        # Container para conteúdo do card com estilo
        content_frame = ttk.Frame(
            card,
            bootstyle="light"  # Alterado para light (branco)
        )
        content_frame.pack(fill=BOTH, expand=YES)
        
        # Título do card
        title_label = ttk.Label(
            content_frame,
            text=title,
            font=("Helvetica", 18, "bold"),
            bootstyle="primary",  # Mantém a cor primária apenas no texto
            justify="center"
        )
        title_label.pack(anchor=CENTER, pady=(0, 10), fill=X)
        
        # Descrição
        desc_label = ttk.Label(
            content_frame,
            text=description,
            wraplength=300,
            bootstyle="secondary",  # Cor secundária para descrição
            justify="center"
        )
        desc_label.pack(anchor=CENTER, pady=(0, 20), fill=X)
        
        # Frame para o botão (para centralizar)
        button_frame = ttk.Frame(
            content_frame,
            bootstyle="light"  # Alterado para light (branco)
        )
        button_frame.pack(side=BOTTOM, fill=X, pady=(10, 0))
        
        # Botão com texto visível (tk.Button com fonte explícita; ttk pode não mostrar texto no Windows)
        _cores = {
            "primary": ("#0d6efd", "white"),
            "info": ("#0dcaf0", "black"),
            "success": ("#198754", "white"),
            "warning": ("#ffc107", "black"),
            "secondary": ("#6c757d", "white"),
            "danger": ("#dc3545", "white"),
        }
        bg, fg = _cores.get(style, ("#0d6efd", "white"))
        access_button = tk.Button(
            button_frame,
            text="ACESSAR",
            font=("Segoe UI", 11, "bold"),
            fg=fg,
            bg=bg,
            activeforeground=fg,
            activebackground=bg,
            relief="flat",
            cursor="hand2",
            command=command,
            padx=30,
            pady=20,
            height=3
        )
        access_button.pack(anchor=CENTER, fill=X, padx=20)

    def publicadores(self):
        # Criar janela de publicadores
        publicadores_window = ttk.Toplevel(self.root)
        publicadores_window.title("Gerenciar Publicadores")
        publicadores_window.geometry("800x600")
        
        # Container principal
        main_container = ttk.Frame(publicadores_window, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        # Título
        title_frame = ttk.Frame(main_container)
        title_frame.pack(fill=X, pady=(0, 20))
        
        ttk.Label(
            title_frame,
            text="Gerenciar Publicadores",
            font=("Helvetica", 20, "bold"),
            bootstyle="primary"
        ).pack(side=LEFT)
        
        # Frame de busca
        search_frame = ttk.Frame(main_container)
        search_frame.pack(fill=X, pady=(0, 10))
        
        search_var = ttk.StringVar()
        
        # Campo de busca
        search_entry = ttk.Entry(
            search_frame,
            textvariable=search_var,
            width=40,
            bootstyle="primary"
        )
        search_entry.pack(side=LEFT, expand=YES, fill=X, padx=(0, 10))
        
        # Frame para filtros de Ancião e Servo Ministerial
        filtros_frame = ttk.LabelFrame(main_container, text="Filtros", padding=10)
        filtros_frame.pack(fill=X, pady=(0, 10))
        
        # Variáveis para os filtros
        filtrar_anciao_var = ttk.BooleanVar()
        filtrar_servo_var = ttk.BooleanVar()
        
        # Checkboxes para filtros
        ttk.Checkbutton(
            filtros_frame,
            text="Ancião",
            variable=filtrar_anciao_var,
            bootstyle="primary-round-toggle"
        ).pack(side=LEFT, padx=(0, 20))
        
        ttk.Checkbutton(
            filtros_frame,
            text="Servo Ministerial",
            variable=filtrar_servo_var,
            bootstyle="primary-round-toggle"
        ).pack(side=LEFT, padx=(0, 20))
        
        # Botão para limpar filtros
        ttk.Button(
            filtros_frame,
            text="Limpar Filtros",
            command=lambda: [filtrar_anciao_var.set(False), filtrar_servo_var.set(False), atualizar_lista()],
            bootstyle="secondary-outline",
            width=15
        ).pack(side=LEFT)
        
        def abrir_modal_adicionar():
            # Criar janela modal
            modal = ttk.Toplevel(publicadores_window)
            modal.title("Adicionar Publicador")
            
            # Container principal
            modal_container = ttk.Frame(modal, padding=20)
            modal_container.pack(fill=BOTH, expand=YES)
            
            # Título
            modal_title_frame = ttk.Frame(modal_container)
            modal_title_frame.pack(fill=X, pady=(0, 20))
            
            ttk.Label(
                modal_title_frame,
                text="Adicionar Publicador",
                font=("Helvetica", 20, "bold"),
                bootstyle="primary"
            ).pack(side=LEFT)
            
            # Frame para os campos
            fields_frame = ttk.Frame(modal_container)
            fields_frame.pack(fill=BOTH, expand=YES, pady=(0, 20))
            
            # Configurar grid para campos
            fields_frame.grid_columnconfigure(1, weight=1)
            
            # Nome
            ttk.Label(
                fields_frame,
                text="Nome:",
                font=("Helvetica", 12),
                bootstyle="secondary"
            ).grid(row=0, column=0, padx=(0, 10), pady=10, sticky="w")
            
            nome_entry = ttk.Entry(
                fields_frame,
                width=40,
                bootstyle="primary"
            )
            nome_entry.grid(row=0, column=1, sticky="ew", pady=10)
            
            # Batizado
            batizado_var = ttk.BooleanVar(value=True)
            ttk.Checkbutton(
                fields_frame,
                text="Batizado",
                variable=batizado_var,
                bootstyle="primary-round-toggle"
            ).grid(row=1, column=0, columnspan=2, pady=10, sticky="w")
            
            # Sexo
            ttk.Label(
                fields_frame,
                text="Sexo:",
                font=("Helvetica", 12),
                bootstyle="secondary"
            ).grid(row=2, column=0, padx=(0, 10), pady=10, sticky="w")
            
            sexo_var = ttk.StringVar(value="Masculino")
            sexo_combo = ttk.Combobox(
                fields_frame,
                textvariable=sexo_var,
                values=["Masculino", "Feminino"],
                width=20,
                state="readonly",
                bootstyle="primary"
            )
            sexo_combo.grid(row=2, column=1, sticky="w", pady=10)
            
            # Frame para permissões
            permissoes_frame = ttk.LabelFrame(fields_frame, text="Permissões", padding=10)
            permissoes_frame.grid(row=3, column=0, columnspan=2, pady=10, sticky="ew")
            
            # Variáveis para permissões
            permissao_escola_var = ttk.BooleanVar(value=True)
            permissao_oracao_var = ttk.BooleanVar(value=True)
            permissao_leitura_var = ttk.BooleanVar(value=True)
            permissao_leitura_sentinela_var = ttk.BooleanVar(value=False)
            permissao_presidente_final_semana_var = ttk.BooleanVar(value=False)
            permissao_audio_video_var = ttk.BooleanVar(value=False)
            permissao_indicador_var = ttk.BooleanVar(value=False)
            permissao_microfone_var = ttk.BooleanVar(value=False)

            # Linha 1: permissões gerais
            add_perm_linha1 = ttk.Frame(permissoes_frame)
            add_perm_linha1.pack(fill=X)

            ttk.Checkbutton(
                add_perm_linha1,
                text="Parte da Escola",
                variable=permissao_escola_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT, padx=(0, 12))

            ttk.Checkbutton(
                add_perm_linha1,
                text="Oração",
                variable=permissao_oracao_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT, padx=(0, 12))

            ttk.Checkbutton(
                add_perm_linha1,
                text="Leitura do Livro",
                variable=permissao_leitura_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT, padx=(0, 12))

            ttk.Checkbutton(
                add_perm_linha1,
                text="Leitura da Sentinela",
                variable=permissao_leitura_sentinela_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT, padx=(0, 12))

            ttk.Checkbutton(
                add_perm_linha1,
                text="Presidente Final de Semana",
                variable=permissao_presidente_final_semana_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT)

            # Linha 2: permissões de salão
            add_perm_linha2 = ttk.Frame(permissoes_frame)
            add_perm_linha2.pack(fill=X, pady=(8, 0))

            ttk.Label(
                add_perm_linha2,
                text="Salão:",
                font=("Helvetica", 10),
                bootstyle="secondary"
            ).pack(side=LEFT, padx=(0, 8))

            ttk.Checkbutton(
                add_perm_linha2,
                text="Áudio/Vídeo",
                variable=permissao_audio_video_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT, padx=(0, 12))

            ttk.Checkbutton(
                add_perm_linha2,
                text="Indicador",
                variable=permissao_indicador_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT, padx=(0, 12))

            ttk.Checkbutton(
                add_perm_linha2,
                text="Microfone",
                variable=permissao_microfone_var,
                bootstyle="primary-round-toggle"
            ).pack(side=LEFT)
            
            # Frame do botão
            button_frame = ttk.Frame(modal_container)
            button_frame.pack(fill=X, pady=(20, 0))
            
            def adicionar_publicador():
                nome = nome_entry.get()
                batizado = batizado_var.get()
                sexo = sexo_var.get()
                permissoes = {
                    "parte_escola": permissao_escola_var.get(),
                    "oracao": permissao_oracao_var.get(),
                    "leitura_livro": permissao_leitura_var.get(),
                    "leitura_sentinela": permissao_leitura_sentinela_var.get(),
                    "presidente_final_semana": permissao_presidente_final_semana_var.get(),
                    "audio_video": permissao_audio_video_var.get(),
                    "indicador": permissao_indicador_var.get(),
                    "microfone": permissao_microfone_var.get()
                }
                if nome.strip():
                    from database import db_ops
                    db_ops.post(nome, batizado, sexo=sexo, permissoes=permissoes)
                    atualizar_lista()
                    modal.destroy()
            
            ttk.Button(
                button_frame,
                text="Adicionar",
                command=adicionar_publicador,
                bootstyle="success",
                padding=(20, 10)
            ).pack(side=LEFT, padx=5)
            
            ttk.Button(
                button_frame,
                text="Cancelar",
                command=modal.destroy,
                bootstyle="secondary",
                padding=(20, 10)
            ).pack(side=LEFT, padx=5)
            
            # Centralizar a janela
            modal.update_idletasks()
            largura_janela = 650
            altura_janela = 520
            x_centralizado = int((modal.winfo_screenwidth() / 2) - (largura_janela / 2))
            y_centralizado = int((modal.winfo_screenheight() / 2) - (altura_janela / 2))
            modal.geometry(f"{largura_janela}x{altura_janela}+{x_centralizado}+{y_centralizado}")
            modal.transient(publicadores_window)
            modal.grab_set()
            
            # Foco no campo de nome
            nome_entry.focus()
        
        # Botão de adicionar publicador
        ttk.Button(
            search_frame,
            text="Adicionar Publicador",
            command=abrir_modal_adicionar,
            bootstyle="primary"
        ).pack(side=RIGHT, padx=(5, 0))
        
        # Botão limpar busca
        clear_button = ttk.Button(
            search_frame,
            text="✕",
            command=lambda: [search_var.set(""), search_entry.insert(0, "Buscar publicador...")],
            bootstyle="secondary-link",
            width=3
        )
        clear_button.pack(side=LEFT, padx=5)
        
        # Frame para a tabela
        table_frame = ttk.Frame(main_container)
        table_frame.pack(fill=BOTH, expand=YES)
        
        # Configurar tabela
        tabela = ttk.Treeview(
            table_frame,
            columns=("nome", "batizado", "anciao", "servo_ministerial", "ultima_parte"),
            show="headings",
            height=15,
            bootstyle="primary"
        )
        
        # Estilo para as linhas alternadas
        tabela.tag_configure('oddrow', background='#f0f0f0')
        tabela.tag_configure('evenrow', background='white')
        
        # Configurar scrollbar
        scrollbar = ttk.Scrollbar(
            table_frame,
            orient="vertical",
            command=tabela.yview,
            bootstyle="primary-round"
        )
        scrollbar.pack(side=RIGHT, fill=Y)
        tabela.configure(yscrollcommand=scrollbar.set)
        
        # Configurar colunas
        tabela.heading("nome", text="Nome", anchor=W)
        tabela.heading("batizado", text="Batizado", anchor=CENTER)
        tabela.heading("anciao", text="Ancião", anchor=CENTER)
        tabela.heading("servo_ministerial", text="Servo Ministerial", anchor=CENTER)
        tabela.heading("ultima_parte", text="Última Parte", anchor=W)
        
        # Configurar largura e alinhamento das colunas
        tabela.column("nome", width=200, anchor=W)
        tabela.column("batizado", width=80, anchor=CENTER)
        tabela.column("anciao", width=80, anchor=CENTER)
        tabela.column("servo_ministerial", width=120, anchor=CENTER)
        tabela.column("ultima_parte", width=200, anchor=W)
        
        tabela.pack(fill=BOTH, expand=YES, padx=(0, 10))
        
        # Frame para botões de ação
        action_frame = ttk.Frame(main_container)
        action_frame.pack(fill=X, pady=(10, 0))
        
        # Variável para armazenar o item selecionado
        selected_item = {'nome': None}
        
        def atualizar_lista():
            try:
                # Limpa a tabela
                for item in tabela.get_children():
                    tabela.delete(item)
                
                # Obtém todos os publicadores
                from database import db_ops
                results = db_ops.getAllPub()
                
                # Filtra baseado na busca por nome
                termo_busca = search_var.get().lower()
                if termo_busca and termo_busca != "buscar publicador...":
                    results = [r for r in results if termo_busca in r["nome"].lower()]
                
                # Filtra por Ancião e/ou Servo Ministerial
                filtrar_anciao = filtrar_anciao_var.get()
                filtrar_servo = filtrar_servo_var.get()
                
                if filtrar_anciao and filtrar_servo:
                    # Se ambos estiverem marcados, mostrar quem é Ancião OU Servo Ministerial
                    results = [r for r in results if r.get("Anciao", False) or r.get("Servo_Ministerial", False)]
                elif filtrar_anciao:
                    # Se apenas Ancião estiver marcado
                    results = [r for r in results if r.get("Anciao", False)]
                elif filtrar_servo:
                    # Se apenas Servo Ministerial estiver marcado
                    results = [r for r in results if r.get("Servo_Ministerial", False)]
                
                # Ordenação
                results.sort(key=lambda x: x["nome"].lower())
                
                # Atualiza a tabela
                for i, result in enumerate(results):
                    row_tags = ('evenrow',) if i % 2 == 0 else ('oddrow',)
                    tabela.insert("", "end", values=(
                        result["nome"],
                        "Sim" if result.get("batizado", False) else "Não",
                        "Sim" if result.get("Anciao", False) else "Não",
                        "Sim" if result.get("Servo_Ministerial", False) else "Não",
                        result.get("ultima_parte", "N/A")
                    ), tags=row_tags)
                
                # Se não houver resultados, mostra uma mensagem na tabela
                if not results:
                    tabela.insert("", "end", values=("Nenhum publicador encontrado", "", "", "", ""))
            
            except Exception as e:
                print(f"Erro ao atualizar lista: {e}")
                tabela.insert("", "end", values=(f"Erro ao carregar publicadores: {str(e)}", "", "", "", ""))
        
        def on_select(event):
            selected = tabela.selection()
            if selected:
                item = tabela.item(selected[0])
                selected_item['nome'] = item['values'][0]
                editar_btn.configure(state="normal")
                excluir_btn.configure(state="disabled")
                excluir_btn.configure(state="normal")
            else:
                selected_item['nome'] = None
                editar_btn.configure(state="disabled")
                excluir_btn.configure(state="disabled")
        
        def editar_selecionado():
            if selected_item['nome']:
                # Criar janela modal
                modal = ttk.Toplevel(publicadores_window)
                modal.title("Editar Publicador")
                
                # Container principal
                edit_container = ttk.Frame(modal, padding=20)
                edit_container.pack(fill=BOTH, expand=YES)
                
                # Título
                title_frame = ttk.Frame(edit_container)
                title_frame.pack(fill=X, pady=(0, 20))
                
                ttk.Label(
                    title_frame,
                    text=f"Editar Publicador - {selected_item['nome']}",
                    font=("Helvetica", 20, "bold"),
                    bootstyle="primary"
                ).pack(side=LEFT)
                
                # Frame para os campos
                fields_frame = ttk.Frame(edit_container)
                fields_frame.pack(fill=BOTH, expand=YES, pady=(0, 20))
                
                # Configurar grid para campos
                fields_frame.grid_columnconfigure(1, weight=1)
                
                # Buscar dados atuais do publicador
                from database import db_ops
                publicador = next((p for p in db_ops.getAllPub() if p["nome"] == selected_item['nome']), None)
                
                if publicador:
                    # Nome
                    ttk.Label(
                        fields_frame,
                        text="Nome:",
                        font=("Helvetica", 12),
                        bootstyle="secondary"
                    ).grid(row=0, column=0, padx=(0, 10), pady=10, sticky="w")
                    
                    nome_entry = ttk.Entry(
                        fields_frame,
                        width=40,
                        bootstyle="primary"
                    )
                    nome_entry.insert(0, publicador["nome"])
                    nome_entry.grid(row=0, column=1, sticky="ew", pady=10)
                    
                    # Batizado
                    batizado_var = ttk.BooleanVar(value=publicador.get("batizado", False))
                    ttk.Checkbutton(
                        fields_frame,
                        text="Batizado",
                        variable=batizado_var,
                        bootstyle="primary-round-toggle"
                    ).grid(row=1, column=0, columnspan=2, pady=10, sticky="w")
                    
                    # Ancião
                    anciao_var = ttk.BooleanVar(value=publicador.get("Anciao", False))
                    ttk.Checkbutton(
                        fields_frame,
                        text="Ancião",
                        variable=anciao_var,
                        bootstyle="primary-round-toggle"
                    ).grid(row=2, column=0, columnspan=2, pady=10, sticky="w")
                    
                    # Servo Ministerial
                    servo_ministerial_var = ttk.BooleanVar(value=publicador.get("Servo_Ministerial", False))
                    ttk.Checkbutton(
                        fields_frame,
                        text="Servo Ministerial",
                        variable=servo_ministerial_var,
                        bootstyle="primary-round-toggle"
                    ).grid(row=3, column=0, columnspan=2, pady=10, sticky="w")
                    
                    # Sexo
                    ttk.Label(
                        fields_frame,
                        text="Sexo:",
                        font=("Helvetica", 12),
                        bootstyle="secondary"
                    ).grid(row=4, column=0, padx=(0, 10), pady=10, sticky="w")
                    
                    sexo_var = ttk.StringVar(value=publicador.get("sexo", "Masculino"))
                    sexo_combo = ttk.Combobox(
                        fields_frame,
                        textvariable=sexo_var,
                        values=["Masculino", "Feminino"],
                        width=20,
                        state="readonly",
                        bootstyle="primary"
                    )
                    sexo_combo.grid(row=4, column=1, sticky="w", pady=10)
                    
                    # Frame para permissões
                    permissoes_frame = ttk.LabelFrame(fields_frame, text="Permissões", padding=10)
                    permissoes_frame.grid(row=5, column=0, columnspan=2, pady=10, sticky="ew")
                    
                    # Obter permissões do publicador ou usar valores padrão
                    permissoes_publicador = publicador.get("permissoes", {
                        "parte_escola": True,
                        "oracao": True,
                        "leitura_livro": True,
                        "leitura_sentinela": False,
                        "presidente_final_semana": False,
                        "audio_video": False,
                        "indicador": False,
                        "microfone": False
                    })

                    # Variáveis para permissões
                    permissao_escola_var = ttk.BooleanVar(value=permissoes_publicador.get("parte_escola", True))
                    permissao_oracao_var = ttk.BooleanVar(value=permissoes_publicador.get("oracao", True))
                    permissao_leitura_var = ttk.BooleanVar(value=permissoes_publicador.get("leitura_livro", True))
                    permissao_leitura_sentinela_var = ttk.BooleanVar(value=permissoes_publicador.get("leitura_sentinela", False))
                    permissao_presidente_final_semana_var = ttk.BooleanVar(value=permissoes_publicador.get("presidente_final_semana", False))
                    permissao_audio_video_var = ttk.BooleanVar(value=permissoes_publicador.get("audio_video", False))
                    permissao_indicador_var = ttk.BooleanVar(value=permissoes_publicador.get("indicador", False))
                    permissao_microfone_var = ttk.BooleanVar(value=permissoes_publicador.get("microfone", False))
                    
                    # Linha 1: permissões gerais
                    perm_linha1 = ttk.Frame(permissoes_frame)
                    perm_linha1.pack(fill=X)

                    ttk.Checkbutton(
                        perm_linha1,
                        text="Parte da Escola",
                        variable=permissao_escola_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT, padx=(0, 12))

                    ttk.Checkbutton(
                        perm_linha1,
                        text="Oração",
                        variable=permissao_oracao_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT, padx=(0, 12))

                    ttk.Checkbutton(
                        perm_linha1,
                        text="Leitura do Livro",
                        variable=permissao_leitura_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT, padx=(0, 12))

                    ttk.Checkbutton(
                        perm_linha1,
                        text="Leitura da Sentinela",
                        variable=permissao_leitura_sentinela_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT, padx=(0, 12))

                    ttk.Checkbutton(
                        perm_linha1,
                        text="Presidente Final de Semana",
                        variable=permissao_presidente_final_semana_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT)

                    # Linha 2: permissões de salão
                    perm_linha2 = ttk.Frame(permissoes_frame)
                    perm_linha2.pack(fill=X, pady=(8, 0))

                    ttk.Label(
                        perm_linha2,
                        text="Salão:",
                        font=("Helvetica", 10),
                        bootstyle="secondary"
                    ).pack(side=LEFT, padx=(0, 8))

                    ttk.Checkbutton(
                        perm_linha2,
                        text="Áudio/Vídeo",
                        variable=permissao_audio_video_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT, padx=(0, 12))

                    ttk.Checkbutton(
                        perm_linha2,
                        text="Indicador",
                        variable=permissao_indicador_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT, padx=(0, 12))

                    ttk.Checkbutton(
                        perm_linha2,
                        text="Microfone",
                        variable=permissao_microfone_var,
                        bootstyle="primary-round-toggle"
                    ).pack(side=LEFT)

                    # Última parte (somente leitura)
                    if "ultima_parte" in publicador:
                        ttk.Label(
                            fields_frame,
                            text="Última Parte:",
                            font=("Helvetica", 12),
                            bootstyle="secondary"
                        ).grid(row=6, column=0, padx=(0, 10), pady=10, sticky="w")
                        
                        ttk.Label(
                            fields_frame,
                            text=publicador.get("ultima_parte", "N/A"),
                            font=("Helvetica", 12),
                            bootstyle="primary"
                        ).grid(row=6, column=1, sticky="w", pady=10)
                    
                    # Botões
                    button_frame = ttk.Frame(edit_container)
                    button_frame.pack(fill=X, pady=(20, 0))
                    
                    def salvar_alteracoes():
                        novo_nome = nome_entry.get()
                        batizado = batizado_var.get()
                        anciao = anciao_var.get()
                        servo_ministerial = servo_ministerial_var.get()
                        sexo = sexo_var.get()
                        permissoes = {
                            "parte_escola": permissao_escola_var.get(),
                            "oracao": permissao_oracao_var.get(),
                            "leitura_livro": permissao_leitura_var.get(),
                            "leitura_sentinela": permissao_leitura_sentinela_var.get(),
                            "presidente_final_semana": permissao_presidente_final_semana_var.get(),
                            "audio_video": permissao_audio_video_var.get(),
                            "indicador": permissao_indicador_var.get(),
                            "microfone": permissao_microfone_var.get()
                        }
                        
                        if novo_nome.strip():
                            from database import db_ops
                            # Se o nome mudou, precisamos excluir o antigo e criar um novo
                            if novo_nome != selected_item['nome']:
                                # Buscar dados do publicador antigo para preservar histórico
                                publicador_antigo = next((p for p in db_ops.getAllPub() if p["nome"] == selected_item['nome']), None)
                                historico_antigo = publicador_antigo.get("historico", []) if publicador_antigo else []
                                ultima_parte_antiga = publicador_antigo.get("ultima_parte", "") if publicador_antigo else ""
                                
                                # Deletar o antigo
                                db_ops.delete(selected_item['nome'])
                                
                                # Criar novo com os dados atualizados
                                db_ops.post(novo_nome, batizado, sexo=sexo, permissoes=permissoes)
                                
                                # Atualizar campos adicionais e histórico
                                db_ops.update_publicador(novo_nome, batizado=batizado, anciao=anciao, servo_ministerial=servo_ministerial)
                                
                                # Restaurar histórico se houver
                                if historico_antigo or ultima_parte_antiga:
                                    try:
                                        if hasattr(db_ops.db, 'find'):
                                            # self.db é a collection de publicadores
                                            collection_publicadores = db_ops.db
                                        else:
                                            # self.db é o database, acessar a collection
                                            collection_publicadores = db_ops.db.database['publicadores']
                                        
                                        collection_publicadores.update_one(
                                            {"nome": novo_nome},
                                            {"$set": {"historico": historico_antigo, "ultima_parte": ultima_parte_antiga}}
                                        )
                                    except Exception as e:
                                        logger.error(f"Erro ao restaurar histórico: {str(e)}")
                            else:
                                # Atualizar todos os campos
                                db_ops.update_publicador(
                                    selected_item['nome'], 
                                    batizado=batizado, 
                                    anciao=anciao, 
                                    servo_ministerial=servo_ministerial,
                                    sexo=sexo,
                                    permissoes=permissoes
                                )
                            
                            atualizar_lista()
                            modal.destroy()
                    
                    ttk.Button(
                        button_frame,
                        text="Salvar",
                        command=salvar_alteracoes,
                        bootstyle="success",
                        padding=(20, 10)
                    ).pack(side=LEFT, padx=5)
                    
                    ttk.Button(
                        button_frame,
                        text="Cancelar",
                        command=modal.destroy,
                        bootstyle="secondary",
                        padding=(20, 10)
                    ).pack(side=LEFT, padx=5)
                    
                    # Centralizar a janela
                    modal.update_idletasks()
                    largura_janela = 650
                    altura_janela = 620
                    x_centralizado = int((modal.winfo_screenwidth() / 2) - (largura_janela / 2))
                    y_centralizado = int((modal.winfo_screenheight() / 2) - (altura_janela / 2))
                    modal.geometry(f"{largura_janela}x{altura_janela}+{x_centralizado}+{y_centralizado}")
                    modal.transient(publicadores_window)
                    modal.grab_set()
                    
                    # Foco no campo de nome
                    nome_entry.focus()
        
        def excluir_selecionado():
            if selected_item['nome']:
                # Criar janela de confirmação customizada
                dialog = ttk.Toplevel(publicadores_window)
                dialog.title("Confirmar Exclusão")
                
                # Tornar a janela modal
                dialog.transient(publicadores_window)
                dialog.grab_set()
                
                # Container principal
                dialog_frame = ttk.Frame(dialog, padding=20)
                dialog_frame.pack(fill=BOTH, expand=YES)
                
                # Mensagem
                ttk.Label(
                    dialog_frame,
                    text=f"Deseja realmente excluir o publicador {selected_item['nome']}?",
                    font=("Helvetica", 12),
                    wraplength=300
                ).pack(pady=(0, 20))
                
                # Frame para botões
                button_frame = ttk.Frame(dialog_frame)
                button_frame.pack(fill=X)
                
                # Variável para armazenar a resposta
                resposta = {'valor': False}
                
                def confirmar():
                    resposta['valor'] = True
                    dialog.destroy()
                
                def cancelar():
                    dialog.destroy()
                
                # Botões
                ttk.Button(
                    button_frame,
                    text="Sim",
                    command=confirmar,
                    bootstyle="danger",
                    width=10
                ).pack(side=LEFT, padx=5)
                
                ttk.Button(
                    button_frame,
                    text="Não",
                    command=cancelar,
                    bootstyle="secondary",
                    width=10
                ).pack(side=LEFT)
                
                # Centralizar a janela
                dialog.update_idletasks()
                width = 400
                height = 150
                x = publicadores_window.winfo_x() + (publicadores_window.winfo_width() // 2) - (width // 2)
                y = publicadores_window.winfo_y() + (publicadores_window.winfo_height() // 2) - (height // 2)
                dialog.geometry(f"{width}x{height}+{x}+{y}")
                
                # Esperar pela resposta
                dialog.wait_window()
                
                # Se confirmou, excluir
                if resposta['valor']:
                    from database import db_ops
                    db_ops.delete(selected_item['nome'])
                    atualizar_lista()
                    selected_item['nome'] = None
                    editar_btn.configure(state="disabled")
                    excluir_btn.configure(state="disabled")
        
        # Binding para seleção
        tabela.bind('<<TreeviewSelect>>', on_select)
        
        # Botões de ação
        editar_btn = ttk.Button(
            action_frame,
            text="Editar",
            command=editar_selecionado,
            bootstyle="primary",
            state="disabled",
            width=15
        )
        editar_btn.pack(side=LEFT, padx=5)
        
        excluir_btn = ttk.Button(
            action_frame,
            text="Excluir",
            command=excluir_selecionado,
            bootstyle="danger",
            state="disabled",
            width=15
        )
        excluir_btn.pack(side=LEFT, padx=5)
        
        # Botão de resetar histórico (separado visualmente)
        def resetar_historico():
            # Criar janela de confirmação
            dialog = ttk.Toplevel(publicadores_window)
            dialog.title("Confirmar Reset de Histórico")
            
            # Tornar a janela modal
            dialog.transient(publicadores_window)
            dialog.grab_set()
            
            # Container principal
            dialog_frame = ttk.Frame(dialog, padding=20)
            dialog_frame.pack(fill=BOTH, expand=YES)
            
            # Mensagem de aviso
            ttk.Label(
                dialog_frame,
                text="⚠️ ATENÇÃO: Esta operação é IRREVERSÍVEL!",
                font=("Helvetica", 14, "bold"),
                bootstyle="danger"
            ).pack(pady=(0, 10))
            
            ttk.Label(
                dialog_frame,
                text="Esta ação irá:\n• Limpar o histórico de TODOS os publicadores\n• Remover TODAS as reuniões cadastradas",
                font=("Helvetica", 11),
                wraplength=400,
                justify="center"
            ).pack(pady=(0, 20))
            
            ttk.Label(
                dialog_frame,
                text="Deseja realmente continuar?",
                font=("Helvetica", 12, "bold"),
                bootstyle="warning"
            ).pack(pady=(0, 20))
            
            # Frame para botões
            button_frame = ttk.Frame(dialog_frame)
            button_frame.pack(fill=X)
            
            # Variável para armazenar a resposta
            resposta = {'valor': False}
            
            def confirmar():
                resposta['valor'] = True
                dialog.destroy()
            
            def cancelar():
                dialog.destroy()
            
            # Botões
            ttk.Button(
                button_frame,
                text="Sim, Resetar Tudo",
                command=confirmar,
                bootstyle="danger",
                width=20
            ).pack(side=LEFT, padx=5)
            
            ttk.Button(
                button_frame,
                text="Cancelar",
                command=cancelar,
                bootstyle="secondary",
                width=15
            ).pack(side=LEFT, padx=5)
            
            # Centralizar a janela
            dialog.update_idletasks()
            width = 450
            height = 250
            x = publicadores_window.winfo_x() + (publicadores_window.winfo_width() // 2) - (width // 2)
            y = publicadores_window.winfo_y() + (publicadores_window.winfo_height() // 2) - (height // 2)
            dialog.geometry(f"{width}x{height}+{x}+{y}")
            
            # Esperar pela resposta
            dialog.wait_window()
            
            # Se confirmou, executar reset
            if resposta['valor']:
                try:
                    from database import db_ops
                    resultado = db_ops.resetar_todo_historico()
                    
                    if resultado['success']:
                        from ttkbootstrap.dialogs import Messagebox
                        Messagebox.show_info(
                            resultado['message'],
                            "Histórico Resetado",
                            parent=publicadores_window
                        )
                        atualizar_lista()
                    else:
                        from ttkbootstrap.dialogs import Messagebox
                        Messagebox.show_error(
                            resultado['message'],
                            "Erro ao Resetar",
                            parent=publicadores_window
                        )
                except Exception as e:
                    from ttkbootstrap.dialogs import Messagebox
                    Messagebox.show_error(
                        f"Erro ao resetar histórico: {str(e)}",
                        "Erro",
                        parent=publicadores_window
                    )
        
        # Separador visual
        ttk.Separator(action_frame, orient="vertical").pack(side=LEFT, fill=Y, padx=10)
        
        resetar_btn = ttk.Button(
            action_frame,
            text="Resetar Todo Histórico",
            command=resetar_historico,
            bootstyle="danger-outline",
            width=20
        )
        resetar_btn.pack(side=LEFT, padx=5)
        
        # Botão de voltar
        voltar_btn = ttk.Button(
            main_container,
            text="Voltar",
            command=publicadores_window.destroy,
            bootstyle="secondary",
            width=15
        )
        voltar_btn.pack(side=BOTTOM, pady=(20, 0))
        
        # Configurar campo de busca e filtros
        search_var.trace("w", lambda *args: atualizar_lista())
        filtrar_anciao_var.trace("w", lambda *args: atualizar_lista())
        filtrar_servo_var.trace("w", lambda *args: atualizar_lista())
        
        # Ícone ou label de busca
        ttk.Label(
            search_frame,
            text="🔍",
            font=("Helvetica", 12),
            bootstyle="secondary"
        ).pack(side=LEFT, padx=(0, 5))
        
        # Texto placeholder usando bind
        search_entry.insert(0, "Buscar publicador...")
        search_entry.bind("<FocusIn>", lambda e: search_entry.delete(0, END) if search_entry.get() == "Buscar publicador..." else None)
        search_entry.bind("<FocusOut>", lambda e: search_entry.insert(0, "Buscar publicador...") if search_entry.get() == "" else None)
        
        # Carregar dados iniciais
        atualizar_lista()
        
        # Centralizar a janela
        publicadores_window.update_idletasks()
        width = publicadores_window.winfo_width()
        height = publicadores_window.winfo_height()
        x = (publicadores_window.winfo_screenwidth() // 2) - (width // 2)
        y = (publicadores_window.winfo_screenheight() // 2) - (height // 2)
        publicadores_window.geometry(f"{width}x{height}+{x}+{y}")

    def historico(self):
        # Configuração da janela
        historico_window = ttk.Toplevel(self.root)
        historico_window.title("Histórico de Reuniões")
        historico_window.geometry("1024x768")
        
        # Container principal
        main_container = ttk.Frame(historico_window, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        # Título
        title_frame = ttk.Frame(main_container)
        title_frame.pack(fill=X, pady=(0, 20))
        
        ttk.Label(
            title_frame,
            text="Histórico de Reuniões",
            font=("Helvetica", 20, "bold"),
            bootstyle="primary"
        ).pack(side=LEFT)
        
        # Frame de busca e filtros
        search_frame = ttk.Frame(main_container)
        search_frame.pack(fill=X, pady=(0, 20))
        
        # Filtro por mês
        ttk.Label(
            search_frame,
            text="Mês:",
            bootstyle="secondary"
        ).pack(side=LEFT, padx=(0, 10))
        
        mes_var = ttk.StringVar()
        meses = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", 
                "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
        mes_combo = ttk.Combobox(
            search_frame,
            textvariable=mes_var,
            values=meses,
            width=15,
            state="readonly",
            bootstyle="primary"
        )
        mes_combo.pack(side=LEFT, padx=(0, 20))
        
        # Filtro por ano
        ttk.Label(
            search_frame,
            text="Ano:",
            bootstyle="secondary"
        ).pack(side=LEFT, padx=(0, 10))
        
        ano_var = ttk.StringVar()
        ano_combo = ttk.Combobox(
            search_frame,
            textvariable=ano_var,
            values=[str(datetime.datetime.now().year - i) for i in range(5)],
            width=10,
            state="readonly",
            bootstyle="primary"
        )
        ano_combo.pack(side=LEFT, padx=(0, 20))
        ano_combo.set(str(datetime.datetime.now().year))
        
        # Botão de buscar
        ttk.Button(
            search_frame,
            text="Buscar",
            command=lambda: carregar_reunioes(),
            bootstyle="primary",
            width=20
        ).pack(side=LEFT, padx=(0, 10))
        
        # Botão de limpar filtros
        ttk.Button(
            search_frame,
            text="Limpar Filtros",
            command=lambda: [mes_var.set(""), ano_var.set(str(datetime.datetime.now().year)), carregar_reunioes()],
            bootstyle="secondary-outline",
            width=20
        ).pack(side=LEFT)
        
        # Frame para a tabela
        table_frame = ttk.Frame(main_container)
        table_frame.pack(fill=BOTH, expand=YES)
        
        # Criar Treeview
        colunas = ("data", "presidente", "semana_ano", "detalhes")
        tree = ttk.Treeview(
            table_frame,
            columns=colunas,
            show="headings",
            bootstyle="primary",
            height=20
        )
        
        # Configurar colunas
        tree.heading("data", text="Data de Criação", anchor=CENTER)
        tree.heading("presidente", text="Presidente", anchor=CENTER)
        tree.heading("semana_ano", text="Semana/Ano", anchor=CENTER)
        tree.heading("detalhes", text="Detalhes", anchor=CENTER)
        
        # Configurar largura das colunas
        tree.column("data", width=300, anchor=CENTER)
        tree.column("presidente", width=300, anchor=CENTER)
        tree.column("semana_ano", width=150, anchor=CENTER)
        tree.column("detalhes", width=100, anchor=CENTER)
        
        # Adicionar scrollbar
        scrollbar = ttk.Scrollbar(
            table_frame,
            orient="vertical",
            command=tree.yview,
            bootstyle="primary-round"
        )
        scrollbar.pack(side=RIGHT, fill=Y)
        tree.configure(yscrollcommand=scrollbar.set)
        tree.pack(fill=BOTH, expand=YES)
        
        def mostrar_detalhes(item):
            # Obter dados da reunião selecionada
            valores = tree.item(item)['values']
            data = valores[0]  # Data formatada
            semana_ano = valores[2]  # "Semana 9-15 DE JUNHO/2025"
            
            # Extrair semana e ano do formato "Semana 9-15 DE JUNHO/2025"
            # Remover "Semana " do início e dividir por "/"
            partes = semana_ano.replace("Semana ", "").split('/')
            semana = partes[0].strip()  # Extrair a semana (ex: "9-15 DE JUNHO")
            ano = int(partes[1])  # Extrair o ano
            
            # Criar janela de detalhes
            detalhes_window = ttk.Toplevel(historico_window)
            detalhes_window.title(f"Detalhes da Reunião - {data}")
            detalhes_window.geometry("600x400")
            
            # Container principal
            main_container = ttk.Frame(detalhes_window, padding=20)
            main_container.pack(fill=BOTH, expand=YES)
            
            # Título
            ttk.Label(
                main_container,
                text=f"Detalhes da Reunião - {data}",
                font=("Helvetica", 16, "bold"),
                bootstyle="primary"
            ).pack(pady=(0, 20))
            
            # Frame para os detalhes
            detalhes_frame = ttk.Frame(main_container)
            detalhes_frame.pack(fill=BOTH, expand=YES)
            
            # Criar Treeview para detalhes
            detalhes_tree = ttk.Treeview(
                detalhes_frame,
                columns=("parte", "participante"),
                show="headings",
                bootstyle="primary",
                height=12
            )
            
            # Configurar colunas
            detalhes_tree.heading("parte", text="Parte", anchor=CENTER)
            detalhes_tree.heading("participante", text="Participante", anchor=CENTER)
            
            # Configurar largura das colunas
            detalhes_tree.column("parte", width=250, anchor=W)
            detalhes_tree.column("participante", width=250, anchor=W)
            
            # Adicionar scrollbar
            scrollbar = ttk.Scrollbar(
                detalhes_frame,
                orient="vertical",
                command=detalhes_tree.yview,
                bootstyle="primary-round"
            )
            scrollbar.pack(side=RIGHT, fill=Y)
            detalhes_tree.configure(yscrollcommand=scrollbar.set)
            detalhes_tree.pack(fill=BOTH, expand=YES, padx=(0, 10))
            
            try:
                from database import db_ops
                # Buscar reunião usando semana e ano
                reuniao = db_ops.buscar_reuniao(ano, semana)
                
                if reuniao:
                    # Lista de todas as partes com seus participantes
                    detalhes = [
                        ("Presidente", reuniao['presidente']),
                        ("Oração Inicial", reuniao['oracao_inicial']),
                        ("Tesouro", reuniao['tesouro']),
                        ("Joias Espirituais", reuniao['joias_espirituais']),
                        ("Leitura da Bíblia", reuniao['leitura_biblia']),
                        ("Escola - Primeira Parte", reuniao['escola']['primeira_parte']),
                        ("Escola - Segunda Parte", reuniao['escola']['segunda_parte']),
                        ("Escola - Terceira Parte", reuniao['escola']['terceira_parte']),
                        ("Escola - Quarta Parte", reuniao['escola']['quarta_parte']),
                        ("Nossa Vida Cristã - Primeira Parte", reuniao['nossa_vida_crista']['primeira_parte']),
                        ("Nossa Vida Cristã - Segunda Parte", reuniao['nossa_vida_crista']['segunda_parte']),
                        ("Estudo de Congregação", reuniao['estudo_congregacao']),
                        ("Oração Final", reuniao['oracao_final'])
                    ]
                    
                    # Filtrar partes que não são "não possui" e inserir na tabela com cores alternadas
                    for i, (parte, participante) in enumerate(detalhes):
                        if participante != "não possui":
                            tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                            detalhes_tree.insert("", END, values=(parte, participante), tags=(tag,))
                    
                    # Configurar cores alternadas
                    detalhes_tree.tag_configure('evenrow', background='#f0f0f0')
                    detalhes_tree.tag_configure('oddrow', background='white')
                else:
                    ttk.Label(
                        main_container,
                        text=f"Nenhum detalhe encontrado para {semana}/{ano}",
                        bootstyle="danger"
                    ).pack(pady=20)
            
            except Exception as e:
                logger.error(f"Erro ao carregar detalhes: {str(e)}")
                ttk.Label(
                    main_container,
                    text=f"Erro ao carregar detalhes: {str(e)}",
                    bootstyle="danger"
                ).pack(pady=20)
            
            # Botão de voltar
            ttk.Button(
                main_container,
                text="Voltar",
                command=detalhes_window.destroy,
                bootstyle="secondary",
                width=15
            ).pack(side=BOTTOM, pady=(20, 0))
            
            # Centralizar a janela
            detalhes_window.update_idletasks()
            width = detalhes_window.winfo_width()
            height = detalhes_window.winfo_height()
            x = (detalhes_window.winfo_screenwidth() // 2) - (width // 2)
            y = (detalhes_window.winfo_screenheight() // 2) - (height // 2)
            detalhes_window.geometry(f"{width}x{height}+{x}+{y}")
            
            # Tornar a janela modal
            detalhes_window.transient(historico_window)
            detalhes_window.grab_set()
        
        def carregar_reunioes():
            # Limpar tabela
            for item in tree.get_children():
                tree.delete(item)
            
            # Obter filtros
            mes = mes_var.get()
            ano = ano_var.get() if ano_var.get() else None
            
            try:
                # Buscar reuniões - quando não há filtros, usar limite alto para retornar todas
                from database import db_ops
                limite = 10000 if not ano and not mes else 100
                reunioes = db_ops.listar_reunioes(
                    ano=int(ano) if ano else None,
                    limite=limite,
                    pagina=1
                )
                
                if not reunioes:
                    tree.insert("", END, values=("Nenhuma reunião encontrada", "", "", ""))
                    return
                
                # Adicionar reuniões na tabela
                for reuniao in reunioes:
                    # Converter data para formato desejado
                    data_reuniao = datetime.datetime.fromisoformat(reuniao['data_reuniao'])
                    semana = reuniao['semana']
                    ano = reuniao['ano']
                    
                    # Calcular datas da semana
                    data_inicio = data_reuniao - datetime.timedelta(days=data_reuniao.weekday())
                    data_fim = data_inicio + datetime.timedelta(days=6)
                    
                    # Mapeamento de meses em português para números
                    meses_pt = {
                        'janeiro': 1, 'fevereiro': 2, 'março': 3, 'abril': 4,
                        'maio': 5, 'junho': 6, 'julho': 7, 'agosto': 8,
                        'setembro': 9, 'outubro': 10, 'novembro': 11, 'dezembro': 12
                    }
                    
                    # Formatar data usando nomes de meses em português
                    nome_mes_inicio = list(meses_pt.keys())[data_inicio.month - 1]
                    nome_mes_fim = list(meses_pt.keys())[data_fim.month - 1]
                    data_formatada = f"{data_inicio.day} de {nome_mes_inicio} a {data_fim.day} de {nome_mes_fim}"
                    
                    # Verificar filtro de mês usando número do mês
                    if mes:
                        mes_filtro_num = meses_pt.get(mes.lower())
                        if mes_filtro_num and data_inicio.month != mes_filtro_num:
                            continue
                    
                    # Inserir na tabela com ícone de seta
                    tree.insert("", END, values=(
                        data_formatada,
                        reuniao['presidente'],
                        f"Semana {semana}/{ano}",
                        "↗"  # Ícone de seta diagonal para cima
                    ))
            
            except Exception as e:
                logger.error(f"Erro ao carregar reuniões: {str(e)}")
                tree.insert("", END, values=(f"Erro ao carregar reuniões: {str(e)}", "", "", ""))
        
        # Configurar evento de clique na linha
        def on_tree_click(event):
            item = tree.identify_row(event.y)
            if item:
                mostrar_detalhes(item)
        
        tree.bind("<ButtonRelease-1>", on_tree_click)
        
        # Função para criar reunião manual (definida dentro do escopo)
        def criar_reuniao_manual_interna():
            # Criar janela modal
            criar_window = ttk.Toplevel(historico_window)
            criar_window.title("Criar Reunião Manual")
            criar_window.geometry("900x800")
            
            # Container principal com scroll
            main_container_criar = ttk.Frame(criar_window, padding=20)
            main_container_criar.pack(fill=BOTH, expand=YES)
            
            # Título
            title_frame_criar = ttk.Frame(main_container_criar)
            title_frame_criar.pack(fill=X, pady=(0, 20))
            
            ttk.Label(
                title_frame_criar,
                text="Criar Reunião Manual",
                font=("Helvetica", 20, "bold"),
                bootstyle="primary"
            ).pack(side=LEFT)
            
            # Frame para seleção de data
            data_frame = ttk.LabelFrame(main_container_criar, text="Seleção de Semana", padding=15)
            data_frame.pack(fill=X, pady=(0, 20))
            
            # Campo para selecionar a segunda-feira da semana
            ttk.Label(
                data_frame,
                text="Data da Segunda-feira:",
                bootstyle="secondary"
            ).pack(side=LEFT, padx=(0, 10))
            
            data_segunda_var = ttk.StringVar()
            data_segunda_entry = ttk.Entry(
                data_frame,
                textvariable=data_segunda_var,
                width=15,
                bootstyle="primary"
            )
            data_segunda_entry.pack(side=LEFT, padx=(0, 10))
            # Calcular próxima segunda-feira
            hoje = datetime.datetime.now()
            dias_para_segunda = (7 - hoje.weekday()) % 7
            if dias_para_segunda == 0 and hoje.weekday() != 0:
                dias_para_segunda = 7
            proxima_segunda = hoje + datetime.timedelta(days=dias_para_segunda)
            data_segunda_entry.insert(0, proxima_segunda.strftime("%d/%m/%Y"))
            
            # Botão para calcular semana
            def calcular_semana():
                try:
                    # Converter data de entrada
                    data_str = data_segunda_var.get().strip()
                    data_segunda = datetime.datetime.strptime(data_str, "%d/%m/%Y")
                    
                    # Garantir que é segunda-feira (ajustar se necessário)
                    dias_para_segunda = (data_segunda.weekday()) % 7
                    if dias_para_segunda != 0:
                        data_segunda = data_segunda - datetime.timedelta(days=dias_para_segunda)
                        data_segunda_var.set(data_segunda.strftime("%d/%m/%Y"))
                    
                    # Calcular domingo (6 dias depois)
                    data_domingo = data_segunda + datetime.timedelta(days=6)
                    
                    # Mapeamento de meses em português
                    meses_pt = {
                        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
                        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
                        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
                    }
                    
                    # Formatar semana no padrão do banco: "9-15 DE JUNHO"
                    if data_segunda.month == data_domingo.month:
                        semana_formatada = f"{data_segunda.day}-{data_domingo.day} DE {meses_pt[data_segunda.month].upper()}"
                    else:
                        # Se a semana cruza meses, usar o mês da segunda-feira
                        semana_formatada = f"{data_segunda.day}-{data_domingo.day} DE {meses_pt[data_segunda.month].upper()}"
                    
                    semana_label_var.set(semana_formatada)
                    
                except Exception as e:
                    from ttkbootstrap.dialogs import Messagebox
                    Messagebox.show_error(f"Erro ao calcular semana: {str(e)}", "Erro", parent=criar_window)
            
            ttk.Button(
                data_frame,
                text="Calcular Semana",
                command=calcular_semana,
                bootstyle="primary-outline"
            ).pack(side=LEFT, padx=(0, 10))
            
            # Label para mostrar a semana formatada
            semana_label_var = ttk.StringVar(value="")
            semana_label = ttk.Label(
                data_frame,
                textvariable=semana_label_var,
                font=("Helvetica", 11, "bold"),
                bootstyle="info"
            )
            semana_label.pack(side=LEFT, padx=(10, 0))
            
            # Calcular semana inicial
            calcular_semana()
            
            # Buscar lista de publicadores
            from database import db_ops
            publicadores = db_ops.getAllPub()
            nomes_publicadores = ["não possui"] + [pub['nome'] for pub in publicadores]
            
            # Frame para campos de partes (com scroll)
            scroll_frame = ttk.Frame(main_container_criar)
            scroll_frame.pack(fill=BOTH, expand=YES)
            
            # Canvas para scroll
            canvas = tk.Canvas(scroll_frame)
            scrollbar_criar = ttk.Scrollbar(scroll_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)
            
            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )
            
            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar_criar.set)
            
            # Frame para partes
            partes_frame = ttk.LabelFrame(scrollable_frame, text="Partes da Reunião", padding=15)
            partes_frame.pack(fill=BOTH, expand=YES, padx=10, pady=10)
            
            # Dicionário para armazenar os comboboxes
            partes_widgets = {}
            
            # Lista de partes na ordem correta
            partes_lista = [
                ("presidente", "Presidente"),
                ("oracao_inicial", "Oração Inicial"),
                ("tesouro", "Tesouro"),
                ("joias_espirituais", "Joias Espirituais"),
                ("leitura_biblia", "Leitura da Bíblia"),
                ("escola_primeira_parte", "Escola - Primeira Parte"),
                ("escola_segunda_parte", "Escola - Segunda Parte"),
                ("escola_terceira_parte", "Escola - Terceira Parte"),
                ("escola_quarta_parte", "Escola - Quarta Parte"),
                ("nvc_primeira_parte", "Nossa Vida Cristã - Primeira Parte"),
                ("nvc_segunda_parte", "Nossa Vida Cristã - Segunda Parte"),
                ("estudo_congregacao", "Estudo de Congregação"),
                ("oracao_final", "Oração Final")
            ]
            
            # Criar campos para cada parte usando Entry + Listbox (similar à tela de criar reunião)
            for i, (key, label) in enumerate(partes_lista):
                row_frame = ttk.Frame(partes_frame)
                row_frame.pack(fill=X, pady=5)
                
                ttk.Label(
                    row_frame,
                    text=f"{label}:",
                    width=30,
                    anchor=W,
                    bootstyle="secondary"
                ).pack(side=LEFT, padx=(0, 10))
                
                # Frame para Entry e Listbox
                input_frame = ttk.Frame(row_frame)
                input_frame.pack(side=LEFT, fill=X, expand=YES)
                
                parte_var = ttk.StringVar(value="não possui")
                parte_entry = ttk.Entry(
                    input_frame,
                    textvariable=parte_var,
                    width=40,
                    bootstyle="primary"
                )
                parte_entry.pack(fill=X, expand=YES)
                
                # Listbox para mostrar opções filtradas
                parte_listbox = tk.Listbox(
                    input_frame,
                    height=5,
                    font=("Helvetica", 10)
                )
                parte_listbox.pack_forget()  # Inicialmente escondido
                
                # Criar classe para manter estado de cada campo
                class CampoPublicador:
                    def __init__(self, entry, listbox, valores_completos):
                        self.entry = entry
                        self.listbox = listbox
                        self.valores_completos = valores_completos
                        self.listbox_visivel = False
                    
                    def update_listbox(self, event=None):
                        texto_atual = self.entry.get().lower()
                        
                        # Se houver "/", pegar apenas o último termo após a barra
                        if '/' in texto_atual:
                            texto_atual = texto_atual.split('/')[-1].strip()
                        
                        # Limpar listbox
                        self.listbox.delete(0, tk.END)
                        
                        # Filtrar valores
                        if texto_atual:
                            valores_filtrados = [v for v in self.valores_completos if texto_atual in v.lower()]
                        else:
                            valores_filtrados = self.valores_completos
                        
                        # Adicionar valores filtrados ao listbox
                        for valor in valores_filtrados:
                            self.listbox.insert(tk.END, valor)
                        
                        # Mostrar/ocultar listbox baseado em se há valores
                        if valores_filtrados and texto_atual:
                            if not self.listbox_visivel:
                                self.listbox.pack(fill=X, expand=YES)
                                self.listbox_visivel = True
                        else:
                            if self.listbox_visivel:
                                self.listbox.pack_forget()
                                self.listbox_visivel = False
                    
                    def on_listbox_select(self, event):
                        selection = self.listbox.curselection()
                        if selection:
                            index = selection[0]
                            selected_name = self.listbox.get(index)
                            current_text = self.entry.get()
                            
                            # Se houver "/", adicionar após a barra
                            if '/' in current_text:
                                parts = current_text.split('/')
                                # Manter tudo antes da última barra e adicionar o novo nome
                                if len(parts) > 1:
                                    # Pegar tudo antes da última barra
                                    antes_barra = '/'.join(parts[:-1])
                                    self.entry.delete(0, tk.END)
                                    self.entry.insert(0, antes_barra + ' / ' + selected_name)
                                else:
                                    self.entry.delete(0, tk.END)
                                    self.entry.insert(0, parts[0] + ' / ' + selected_name)
                            else:
                                # Se não houver "/", substituir
                                self.entry.delete(0, tk.END)
                                self.entry.insert(0, selected_name)
                            
                            # Ocultar listbox após seleção
                            self.listbox.pack_forget()
                            self.listbox_visivel = False
                            # Focar no entry novamente
                            self.entry.focus_set()
                    
                    def on_focus_out(self, event):
                        # Ocultar listbox quando perder foco
                        if self.listbox_visivel:
                            self.listbox.pack_forget()
                            self.listbox_visivel = False
                
                # Criar instância do campo
                campo = CampoPublicador(parte_entry, parte_listbox, nomes_publicadores)
                
                # Bind eventos
                parte_entry.bind('<KeyRelease>', campo.update_listbox)
                parte_listbox.bind('<<ListboxSelect>>', campo.on_listbox_select)
                parte_listbox.bind('<Double-Button-1>', campo.on_listbox_select)
                parte_entry.bind('<FocusOut>', campo.on_focus_out)
                
                partes_widgets[key] = parte_var
            
            canvas.pack(side=LEFT, fill=BOTH, expand=YES)
            scrollbar_criar.pack(side=RIGHT, fill=Y)
            
            # Frame para botões
            button_frame_criar = ttk.Frame(main_container_criar)
            button_frame_criar.pack(fill=X, pady=(20, 0))
            
            def salvar_reuniao():
                try:
                    # Validar data
                    data_str = data_segunda_var.get().strip()
                    data_segunda = datetime.datetime.strptime(data_str, "%d/%m/%Y")
                    
                    # Garantir que é segunda-feira
                    dias_para_segunda = (data_segunda.weekday()) % 7
                    if dias_para_segunda != 0:
                        data_segunda = data_segunda - datetime.timedelta(days=dias_para_segunda)
                        data_segunda_var.set(data_segunda.strftime("%d/%m/%Y"))
                    
                    # Calcular domingo
                    data_domingo = data_segunda + datetime.timedelta(days=6)
                    
                    # Formatar semana
                    meses_pt = {
                        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
                        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
                        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
                    }
                    
                    # Formatar semana no padrão do banco: "9-15 DE JUNHO"
                    if data_segunda.month == data_domingo.month:
                        semana_formatada = f"{data_segunda.day}-{data_domingo.day} DE {meses_pt[data_segunda.month].upper()}"
                    else:
                        # Se a semana cruza meses, usar o mês da segunda-feira
                        semana_formatada = f"{data_segunda.day}-{data_domingo.day} DE {meses_pt[data_segunda.month].upper()}"
                    
                    semana = semana_formatada
                    ano = data_segunda.year
                    
                    # Preparar dados da reunião
                    dados_reuniao = {
                        'ano': ano,
                        'semana': semana,
                        'data_reuniao': data_segunda.isoformat(),  # Data da segunda-feira
                        'presidente': partes_widgets['presidente'].get().strip() or "não possui",
                        'oracao_inicial': partes_widgets['oracao_inicial'].get().strip() or "não possui",
                        'tesouro': partes_widgets['tesouro'].get().strip() or "não possui",
                        'joias_espirituais': partes_widgets['joias_espirituais'].get().strip() or "não possui",
                        'leitura_biblia': partes_widgets['leitura_biblia'].get().strip() or "não possui",
                        'escola': {
                            'primeira_parte': partes_widgets['escola_primeira_parte'].get().strip() or "não possui",
                            'segunda_parte': partes_widgets['escola_segunda_parte'].get().strip() or "não possui",
                            'terceira_parte': partes_widgets['escola_terceira_parte'].get().strip() or "não possui",
                            'quarta_parte': partes_widgets['escola_quarta_parte'].get().strip() or "não possui"
                        },
                        'nossa_vida_crista': {
                            'primeira_parte': partes_widgets['nvc_primeira_parte'].get().strip() or "não possui",
                            'segunda_parte': partes_widgets['nvc_segunda_parte'].get().strip() or "não possui"
                        },
                        'estudo_congregacao': partes_widgets['estudo_congregacao'].get().strip() or "não possui",
                        'oracao_final': partes_widgets['oracao_final'].get().strip() or "não possui"
                    }
                    
                    # Salvar reunião
                    resultado = db_ops.salvar_reuniao(dados_reuniao)
                    
                    if resultado['success']:
                        # Recarregar lista de reuniões antes de fechar a janela
                        carregar_reunioes()
                        # Mostrar mensagem de sucesso usando a janela pai
                        from ttkbootstrap.dialogs import Messagebox
                        Messagebox.show_info(
                            f"Reunião criada com sucesso!\nSemana: {semana}\nAno: {ano}",
                            "Sucesso",
                            parent=historico_window
                        )
                        criar_window.destroy()
                    else:
                        from ttkbootstrap.dialogs import Messagebox
                        Messagebox.show_error(
                            resultado['message'],
                            "Erro ao Salvar",
                            parent=historico_window
                        )
                        
                except ValueError as ve:
                    from ttkbootstrap.dialogs import Messagebox
                    Messagebox.show_error(
                        f"Data inválida. Use o formato DD/MM/AAAA",
                        "Erro de Validação",
                        parent=historico_window
                    )
                except Exception as e:
                    from ttkbootstrap.dialogs import Messagebox
                    Messagebox.show_error(
                        f"Erro ao salvar reunião: {str(e)}",
                        "Erro",
                        parent=historico_window
                    )
            
            ttk.Button(
                button_frame_criar,
                text="Salvar Reunião",
                command=salvar_reuniao,
                bootstyle="success",
                width=20
            ).pack(side=LEFT, padx=5)
            
            ttk.Button(
                button_frame_criar,
                text="Cancelar",
                command=criar_window.destroy,
                bootstyle="secondary",
                width=15
            ).pack(side=LEFT, padx=5)
            
            # Centralizar janela
            criar_window.update_idletasks()
            width = criar_window.winfo_width()
            height = criar_window.winfo_height()
            x = (criar_window.winfo_screenwidth() // 2) - (width // 2)
            y = (criar_window.winfo_screenheight() // 2) - (height // 2)
            criar_window.geometry(f"{width}x{height}+{x}+{y}")
            criar_window.transient(historico_window)
            criar_window.grab_set()
            
            # Foco no campo de data
            data_segunda_entry.focus()
        
        # Botão de criar reunião
        ttk.Button(
            title_frame,
            text="Criar Reunião",
            command=criar_reuniao_manual_interna,
            bootstyle="success",
            width=20
        ).pack(side=RIGHT, padx=(10, 0))
        
        # Botão de voltar
        voltar_btn = ttk.Button(
            main_container,
            text="Voltar",
            command=historico_window.destroy,
            bootstyle="secondary",
            width=15
        )
        voltar_btn.pack(side=BOTTOM, pady=(20, 0))
        
        # Carregar dados iniciais
        carregar_reunioes()
        
        # Centralizar a janela
        historico_window.update_idletasks()
        width = historico_window.winfo_width()
        height = historico_window.winfo_height()
        x = (historico_window.winfo_screenwidth() // 2) - (width // 2)
        y = (historico_window.winfo_screenheight() // 2) - (height // 2)
        historico_window.geometry(f"{width}x{height}+{x}+{y}")

    def criar_quadro_de_anuncio(self):
        quadro_de_anuncio = ttk.Toplevel(self.root)
        quadro_de_anuncio.title("Criar Reunião")
        
        # Container principal
        main_container = ttk.Frame(quadro_de_anuncio, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        # Título
        title_frame = ttk.Frame(main_container)
        title_frame.pack(fill=X, pady=(0, 20))
        
        ttk.Label(
            title_frame,
            text="Criar Reunião de Meio de Semana",
            font=("Helvetica", 20, "bold"),
            bootstyle="primary"
        ).pack(side=LEFT)
        
        # Frame para os campos
        fields_frame = ttk.Frame(main_container)
        fields_frame.pack(fill=X, pady=(0, 20))
        
        # Configurar grid para campos
        fields_frame.grid_columnconfigure(1, weight=1)
        
        # URL
        ttk.Label(
            fields_frame,
            text="URL:",
            font=("Helvetica", 12),
            bootstyle="secondary"
        ).grid(row=0, column=0, padx=(0, 10), pady=10, sticky="w")
        
        entry_url = ttk.Entry(
            fields_frame,
            width=50,
            bootstyle="primary"
        )
        entry_url.grid(row=0, column=1, sticky="ew", pady=10)
        
        # Quantidade de Semanas
        ttk.Label(
            fields_frame,
            text="Quantidade de Semanas:",
            font=("Helvetica", 12),
            bootstyle="secondary"
        ).grid(row=1, column=0, padx=(0, 10), pady=10, sticky="w")
        
        entry_qtdSemanas = ttk.Entry(
            fields_frame,
            width=10,
            bootstyle="primary"
        )
        entry_qtdSemanas.grid(row=1, column=1, sticky="w", pady=10)
        
        # Nome do Arquivo
        ttk.Label(
            fields_frame,
            text="Nome do Arquivo:",
            font=("Helvetica", 12),
            bootstyle="secondary"
        ).grid(row=2, column=0, padx=(0, 10), pady=10, sticky="w")
        
        entry_nome = ttk.Entry(
            fields_frame,
            width=50,
            bootstyle="primary"
        )
        entry_nome.grid(row=2, column=1, sticky="ew", pady=10)
        
        # Idioma
        ttk.Label(
            fields_frame,
            text="Idioma:",
            font=("Helvetica", 12),
            bootstyle="secondary"
        ).grid(row=3, column=0, padx=(0, 10), pady=10, sticky="w")
        
        variavel = ttk.StringVar(quadro_de_anuncio)
        variavel.set("pt")
        
        idioma_combo = ttk.Combobox(
            fields_frame,
            textvariable=variavel,
            values=["pt"],
            state="readonly",
            width=10,
            bootstyle="primary"
        )
        idioma_combo.grid(row=3, column=1, sticky="w", pady=10)
        
        # Checkbox para base de publicadores
        utilizar_base_var = ttk.BooleanVar()
        ttk.Checkbutton(
            fields_frame,
            text="Utilizar base de publicadores",
            variable=utilizar_base_var,
            bootstyle="primary-round-toggle"
        ).grid(row=4, column=0, columnspan=2, pady=10, sticky="w")
        
        # Checkbox para gerar com publicadores automaticamente
        gerar_com_publicadores_var = ttk.BooleanVar()
        ttk.Checkbutton(
            fields_frame,
            text="Gerar com Publicadores (Seleção Automática)",
            variable=gerar_com_publicadores_var,
            bootstyle="primary-round-toggle"
        ).grid(row=5, column=0, columnspan=2, pady=10, sticky="w")
        
        # Botão de enviar
        button_frame = ttk.Frame(main_container)
        button_frame.pack(fill=X, pady=(20, 0))
        
        gerar_btn = ttk.Button(
            button_frame,
            text="Gerar Reunião",
            bootstyle="success",
            padding=(20, 15)
        )
        gerar_btn.pack(anchor=CENTER)
        
        # Frame para loading (sempre presente, mas inicialmente vazio)
        loading_frame = ttk.Frame(main_container)
        # Criar elementos do loading (inicialmente não empacotados)
        loading_label = ttk.Label(
            loading_frame,
            text="Gerando documento... Por favor, aguarde.",
            font=("Helvetica", 14, "bold"),
            bootstyle="info"
        )
        
        progress_var = tk.DoubleVar(value=0)
        loading_progress = ttk.Progressbar(
            loading_frame,
            mode='determinate',
            bootstyle="info-striped",
            length=400,
            maximum=100,
            variable=progress_var
        )
        status_label = ttk.Label(
            loading_frame,
            text="",
            font=("Helvetica", 10),
            bootstyle="secondary"
        )
        
        def mostrar_loading():
            gerar_btn.configure(state="disabled")
            button_frame.pack_forget()
            loading_label.pack(pady=(15, 5))
            progress_var.set(0)
            loading_progress.pack(pady=(0, 5))
            status_label.configure(text="Iniciando...")
            status_label.pack(pady=(0, 15))
            loading_frame.pack(fill=X, pady=(20, 0))
            button_frame.pack(fill=X, pady=(10, 0))
            quadro_de_anuncio.update_idletasks()
            quadro_de_anuncio.update()

        def esconder_loading():
            loading_label.pack_forget()
            loading_progress.pack_forget()
            status_label.pack_forget()
            loading_frame.pack_forget()
            gerar_btn.configure(state="normal")
            quadro_de_anuncio.update_idletasks()
            quadro_de_anuncio.update()
        
        def processar_geracao(url, nome_arquivo, idioma, utilizarBase, gerarComPublicadores, qtdSemanas):
            """Executa a geração em thread separada"""
            def atualizar_progresso(texto, valor):
                def _update(t=texto, v=valor):
                    progress_var.set(v)
                    status_label.configure(text=t)
                    quadro_de_anuncio.update_idletasks()
                quadro_de_anuncio.after(0, _update)

            try:
                print(url, nome_arquivo, idioma, utilizarBase, gerarComPublicadores)

                s140.s140.gerar_s140(url, nome_arquivo, idioma, utilizarBase, gerarComPublicadores, qtdSemanas, progress_cb=atualizar_progresso)
                
                # Esconder loading e mostrar sucesso
                quadro_de_anuncio.after(0, esconder_loading)
                quadro_de_anuncio.after(0, lambda: Messagebox.show_info(
                    "Documento gerado com sucesso!",
                    "Sucesso"
                ))
                
            except Exception as e:
                # Esconder loading e mostrar erro
                erro_msg = str(e)
                quadro_de_anuncio.after(0, esconder_loading)
                quadro_de_anuncio.after(0, lambda: Messagebox.show_error(
                    f"Erro ao gerar documento:\n{erro_msg}",
                    "Erro"
                ))
                logger.error(f"Erro ao gerar reunião: {str(e)}")
        
        def enviar():
            """Inicia o processo de geração em thread separada"""
            # Validar campos obrigatórios
            if not entry_url.get().strip():
                Messagebox.show_warning("Por favor, preencha a URL.", "Campo Obrigatório")
                return
            if not entry_nome.get().strip():
                Messagebox.show_warning("Por favor, preencha o nome do arquivo.", "Campo Obrigatório")
                return
            try:
                qtd = int(entry_qtdSemanas.get())
                if qtd <= 0:
                    Messagebox.show_warning("A quantidade de semanas deve ser maior que zero.", "Valor Inválido")
                    return
            except ValueError:
                Messagebox.show_warning("Por favor, insira um número válido para a quantidade de semanas.", "Valor Inválido")
                return
            
            # Pegar os dados antes de iniciar a thread
            url = entry_url.get()
            nome_arquivo = entry_nome.get()
            idioma = variavel.get()
            utilizarBase = utilizar_base_var.get()
            gerarComPublicadores = gerar_com_publicadores_var.get()
            qtdSemanas = int(entry_qtdSemanas.get())
            
            # Validar: se gerar com publicadores, deve estar marcado "utilizar base"
            if gerarComPublicadores and not utilizarBase:
                Messagebox.show_warning(
                    "Para usar 'Gerar com Publicadores', é necessário marcar 'Utilizar base de publicadores'.",
                    "Validação"
                )
                return
            
            # Mostrar loading imediatamente (antes de iniciar a thread)
            print("Mostrando loading...")
            mostrar_loading()
            print("Loading mostrado")
            
            # Executar em thread separada
            thread = threading.Thread(target=processar_geracao, args=(url, nome_arquivo, idioma, utilizarBase, gerarComPublicadores, qtdSemanas), daemon=True)
            thread.start()
            print("Thread iniciada")
        
        gerar_btn.configure(command=enviar)
        
        # Botão de voltar
        voltar_btn = ttk.Button(
            main_container,
            text="Voltar",
            command=quadro_de_anuncio.destroy,
            bootstyle="secondary",
            width=15
        )
        voltar_btn.pack(side=BOTTOM, pady=(20, 0))
        
        # Centralizar a janela sem travar tamanho (janela cresce ao exibir loading)
        quadro_de_anuncio.update_idletasks()
        width = quadro_de_anuncio.winfo_reqwidth()
        height = quadro_de_anuncio.winfo_reqheight()
        x = (quadro_de_anuncio.winfo_screenwidth() // 2) - (width // 2)
        y = (quadro_de_anuncio.winfo_screenheight() // 2) - (height // 2)
        quadro_de_anuncio.geometry(f"+{x}+{y}")

    def criar_reuniao_final_semana(self):
        """Janela para criar reunião de final de semana."""
        janela = ttk.Toplevel(self.root)
        janela.title("Criar Reunião Final de Semana")
        janela.geometry("600x420")
        janela.grab_set()

        main_container = ttk.Frame(janela, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        ttk.Label(main_container, text="Criar Reunião Final de Semana", font=("Helvetica", 20, "bold"), bootstyle="primary").pack(pady=(0, 20))
        
        fields = ttk.Frame(main_container)
        fields.pack(fill=BOTH, expand=YES, pady=(0, 20))
        fields.grid_columnconfigure(1, weight=1)
        
        ttk.Label(fields, text="URL:", font=("Helvetica", 12), bootstyle="secondary").grid(row=0, column=0, padx=(0, 10), pady=10, sticky="w")
        entry_url = ttk.Entry(fields, width=50, bootstyle="primary")
        entry_url.grid(row=0, column=1, sticky="ew", pady=10)
        entry_url.insert(0, "https://wol.jw.org/pt/wol/meetings/r5/lp-t/2026/6")
        
        ttk.Label(fields, text="Nome do Arquivo:", font=("Helvetica", 12), bootstyle="secondary").grid(row=1, column=0, padx=(0, 10), pady=10, sticky="w")
        entry_nome = ttk.Entry(fields, width=50, bootstyle="primary")
        entry_nome.grid(row=1, column=1, sticky="ew", pady=10)
        
        ttk.Label(fields, text="Idioma:", font=("Helvetica", 12), bootstyle="secondary").grid(row=2, column=0, padx=(0, 10), pady=10, sticky="w")
        idioma_var = ttk.StringVar(value="pt")
        ttk.Combobox(fields, textvariable=idioma_var, values=["pt"], state="readonly", width=10, bootstyle="primary").grid(row=2, column=1, sticky="w", pady=10)
        
        auto_presidente_var = ttk.BooleanVar()
        auto_leitor_var = ttk.BooleanVar()
        ttk.Checkbutton(fields, text="Seleção automática do Presidente", variable=auto_presidente_var, bootstyle="primary-round-toggle").grid(row=3, column=0, columnspan=2, pady=5, sticky="w")
        ttk.Checkbutton(fields, text="Seleção automática do Leitor da Sentinela", variable=auto_leitor_var, bootstyle="primary-round-toggle").grid(row=4, column=0, columnspan=2, pady=5, sticky="w")
        
        loading_frame = ttk.Frame(main_container)
        loading_label = ttk.Label(loading_frame, text="Gerando documento...", font=("Helvetica", 14, "bold"), bootstyle="info")
        loading_progress = ttk.Progressbar(loading_frame, mode='indeterminate', bootstyle="info-striped", length=400)
        
        def processar():
            def cleanup():
                loading_progress.stop()
                loading_label.pack_forget()
                loading_progress.pack_forget()
                loading_frame.pack_forget()
                gerar_btn.configure(state="normal")

            def continuar_na_main_thread(semanas):
                """Executa modal dirigente, modal dados, geração dos dois docs e salvamento na thread principal."""
                try:
                    dirigente = final_semana.FinalSemana.solicitar_dirigente_sentinela(janela)
                    if dirigente is None:
                        cleanup()
                        return
                    if not final_semana.FinalSemana.solicitar_dados_usuario(semanas, parent=janela):
                        cleanup()
                        return
                    nome_base = entry_nome.get().strip()
                    final_semana.FinalSemana.criar_documento_sentinela(semanas, nome_base, dirigente)
                    final_semana.FinalSemana.criar_documento_oradores(semanas, nome_base)
                    final_semana.FinalSemana.salvar_historico_final_semana(semanas, nome_base, dirigente)
                    Messagebox.show_info("Documentos gerados com sucesso!", "Sucesso")
                except Exception as e:
                    Messagebox.show_error(f"Erro:\n{str(e)}", "Erro", parent=janela)
                    logger.error(f"Erro ao gerar final de semana: {str(e)}")
                finally:
                    cleanup()

            try:
                # Fase 1 (thread): scraping e seleção - sem GUI (9 semanas)
                semanas = final_semana.FinalSemana.buscar_titulos_meetings(entry_url.get().strip(), 9)
                if not semanas:
                    raise ValueError("Não foi possível extrair os títulos da Sentinela. Verifique a URL.")
                if auto_presidente_var.get():
                    final_semana.FinalSemana.selecionar_presidente_automaticamente(semanas)
                if auto_leitor_var.get():
                    final_semana.FinalSemana.selecionar_leitor_automaticamente(semanas)
                janela.after(0, lambda s=semanas: continuar_na_main_thread(s))
            except Exception as e:
                err_msg = str(e)
                janela.after(0, lambda m=err_msg: Messagebox.show_error(f"Erro:\n{m}", "Erro", parent=janela))
                logger.error(f"Erro ao gerar final de semana: {err_msg}")
                janela.after(0, cleanup)
        
        def enviar():
            if not entry_url.get().strip():
                Messagebox.show_warning("Preencha a URL.", "Campo Obrigatório", parent=janela)
                return
            if not entry_nome.get().strip():
                Messagebox.show_warning("Preencha o nome do arquivo.", "Campo Obrigatório", parent=janela)
                return
            gerar_btn.configure(state="disabled")
            loading_label.pack(pady=(15, 10))
            loading_progress.pack(pady=(0, 15))
            loading_frame.pack(fill=X, pady=(20, 0))
            loading_progress.start(10)
            janela.update_idletasks()
            threading.Thread(target=processar, daemon=True).start()
        
        gerar_btn = ttk.Button(main_container, text="Gerar Reunião", bootstyle="success", padding=(20, 15), command=enviar)
        gerar_btn.pack(pady=(20, 0))
        
        ttk.Button(main_container, text="Voltar", command=janela.destroy, bootstyle="secondary", width=15).pack(side=BOTTOM, pady=(20, 0))
        
        janela.update_idletasks()
        w, h = 600, 400
        x = (janela.winfo_screenwidth() // 2) - (w // 2)
        y = (janela.winfo_screenheight() // 2) - (h // 2)
        janela.geometry(f"{w}x{h}+{x}+{y}")

    def historico_final_semana(self):
        """Exibe histórico de reuniões de final de semana."""
        hist_window = ttk.Toplevel(self.root)
        hist_window.title("Histórico Final de Semana")
        hist_window.geometry("900x600")
        
        main = ttk.Frame(hist_window, padding=20)
        main.pack(fill=BOTH, expand=YES)
        
        ttk.Label(main, text="Histórico de Reuniões Final de Semana", font=("Helvetica", 20, "bold"), bootstyle="primary").pack(pady=(0, 20))
        
        filtro_frame = ttk.Frame(main)
        filtro_frame.pack(fill=X, pady=(0, 10))
        ttk.Label(filtro_frame, text="Ano:", bootstyle="secondary").pack(side=LEFT, padx=(0, 5))
        ano_var = ttk.StringVar(value=str(datetime.datetime.now().year))
        ano_combo = ttk.Combobox(filtro_frame, textvariable=ano_var, width=8, values=[str(y) for y in range(2020, 2031)])
        ano_combo.pack(side=LEFT, padx=(0, 15))
        ttk.Label(filtro_frame, text="Mês:", bootstyle="secondary").pack(side=LEFT, padx=(0, 5))
        mes_var = ttk.StringVar(value="")
        mes_combo = ttk.Combobox(filtro_frame, textvariable=mes_var, width=10, values=["", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12"])
        mes_combo.pack(side=LEFT, padx=(0, 10))
        
        tree_frame = ttk.Frame(main)
        tree_frame.pack(fill=BOTH, expand=YES, pady=(10, 0))
        columns = ("ano", "mes", "arquivo", "data")
        tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=15)
        tree.heading("ano", text="Ano")
        tree.heading("mes", text="Mês")
        tree.heading("arquivo", text="Arquivo")
        tree.heading("data", text="Data Criação")
        tree.column("ano", width=80)
        tree.column("mes", width=80)
        tree.column("arquivo", width=300)
        tree.column("data", width=180)
        tree.pack(side=LEFT, fill=BOTH, expand=YES)
        scroll = ttk.Scrollbar(tree_frame, orient=VERTICAL, command=tree.yview)
        scroll.pack(side=RIGHT, fill=Y)
        tree.configure(yscrollcommand=scroll.set)
        
        def mostrar_detalhes_final_semana(item):
            valores = tree.item(item)['values']
            if len(valores) < 2:
                return
            try:
                ano = int(valores[0])
                mes = int(valores[1])
            except (ValueError, TypeError):
                return
            reuniao = buscar_reuniao_final_semana(ano, mes)
            detalhes_window = ttk.Toplevel(hist_window)
            detalhes_window.title(f"Detalhes da Reunião - {valores[2]} ({ano}/{mes})")
            detalhes_window.geometry("650x450")
            det_main = ttk.Frame(detalhes_window, padding=20)
            det_main.pack(fill=BOTH, expand=YES)
            ttk.Label(det_main, text=f"Detalhes da Reunião - {valores[2]}", font=("Helvetica", 16, "bold"), bootstyle="primary").pack(pady=(0, 20))
            det_frame = ttk.Frame(det_main)
            det_frame.pack(fill=BOTH, expand=YES)
            det_tree = ttk.Treeview(det_frame, columns=("parte", "participante"), show="headings", bootstyle="primary", height=15)
            det_tree.heading("parte", text="Parte", anchor=CENTER)
            det_tree.heading("participante", text="Participante", anchor=CENTER)
            det_tree.column("parte", width=250, anchor=W)
            det_tree.column("participante", width=300, anchor=W)
            det_scroll = ttk.Scrollbar(det_frame, orient=VERTICAL, command=det_tree.yview, bootstyle="primary-round")
            det_scroll.pack(side=RIGHT, fill=Y)
            det_tree.configure(yscrollcommand=det_scroll.set)
            det_tree.pack(fill=BOTH, expand=YES, padx=(0, 10))
            if reuniao:
                if reuniao.get('dirigente'):
                    det_tree.insert("", END, values=("Dirigente de Sentinela", reuniao.get('dirigente')), tags=('evenrow',))
                if reuniao.get('semanas'):
                    meses_pt = {
                        1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
                        5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
                        9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
                    }
                    nome_mes = meses_pt.get(reuniao.get('mes'), str(reuniao.get('mes', '')))
                    ano_reuniao = reuniao.get('ano', '')
                    for i, semana in enumerate(reuniao['semanas']):
                        campo_semana = (semana.get('semana') or '').strip()
                        if campo_semana and len(campo_semana) <= 30 and " de " in campo_semana:
                            rotulo = campo_semana
                        else:
                            rotulo = f"Semana {i+1} de {nome_mes} de {ano_reuniao}"
                        detalhes = [
                            (f"{rotulo} - Título Estudo", semana.get('titulo_estudo', '')),
                            (f"{rotulo} - Tema Discurso", semana.get('tema_discurso', '')),
                            (f"{rotulo} - Orador", semana.get('orador', '')),
                            (f"{rotulo} - Presidente", semana.get('presidente', '')),
                            (f"{rotulo} - Leitor Sentinela", semana.get('leitor_sentinela', ''))
                        ]
                        for j, (parte, participante) in enumerate(detalhes):
                            if participante:
                                tag = 'evenrow' if (i * 5 + j) % 2 == 0 else 'oddrow'
                                det_tree.insert("", END, values=(parte, participante), tags=(tag,))
                    det_tree.tag_configure('evenrow', background='#f0f0f0')
                    det_tree.tag_configure('oddrow', background='white')
            if not reuniao or not reuniao.get('semanas'):
                ttk.Label(det_main, text="Nenhum detalhe encontrado para esta reunião.", bootstyle="danger").pack(pady=20)

            btn_frame = ttk.Frame(det_main)
            btn_frame.pack(side=BOTTOM, pady=(20, 0))

            def ao_excluir():
                r = Messagebox.yesno("Excluir esta reunião do histórico? As participações (Leitura Sentinela e Presidente) serão removidas do histórico de cada publicador.", title="Confirmar exclusão")
                if r != "Yes":
                    return
                resultado = excluir_reuniao_final_semana(ano, mes)
                if resultado.get("success"):
                    Messagebox.ok("Reunião excluída. O histórico dos publicadores foi atualizado.", title="Excluído")
                    detalhes_window.destroy()
                    carregar()
                else:
                    Messagebox.show_error(resultado.get("message", "Erro ao excluir."), title="Erro")

            def ao_editar():
                import copy
                semanas_copy = copy.deepcopy(reuniao.get("semanas", []))
                if not final_semana.FinalSemana.solicitar_dados_usuario(semanas_copy, parent=detalhes_window):
                    return
                excluir_reuniao_final_semana(ano, mes)
                nome_arquivo = reuniao.get("nome_arquivo", "")
                dirigente = reuniao.get("dirigente", "")
                dados = {
                    "ano": ano,
                    "mes": mes,
                    "nome_arquivo": nome_arquivo,
                    "dirigente": dirigente,
                    "semanas": semanas_copy,
                    "data_criacao": reuniao.get("data_criacao", ""),
                }
                db_ops.salvar_reuniao_final_semana(dados)
                try:
                    final_semana.FinalSemana.criar_documento_sentinela(semanas_copy, nome_arquivo, dirigente)
                    final_semana.FinalSemana.criar_documento_oradores(semanas_copy, nome_arquivo)
                except Exception as ex:
                    logger.exception("Erro ao gerar documentos na edição")
                    Messagebox.show_warning(f"Reunião atualizada no banco, mas os documentos Word não foram gerados: {ex}", title="Aviso")
                else:
                    Messagebox.ok("Reunião atualizada. Documentos Sentinela e Oradores gerados em documentosCriados/.", title="Atualizado")
                detalhes_window.destroy()
                carregar()

            ttk.Button(btn_frame, text="Editar", command=ao_editar, bootstyle="info", width=12).pack(side=LEFT, padx=(0, 10))
            ttk.Button(btn_frame, text="Excluir", command=ao_excluir, bootstyle="danger", width=12).pack(side=LEFT, padx=(0, 10))
            ttk.Button(btn_frame, text="Voltar", command=detalhes_window.destroy, bootstyle="secondary", width=12).pack(side=LEFT)
            detalhes_window.update_idletasks()
            w, h = 650, 450
            x = (detalhes_window.winfo_screenwidth() // 2) - (w // 2)
            y = (detalhes_window.winfo_screenheight() // 2) - (h // 2)
            detalhes_window.geometry(f"{w}x{h}+{x}+{y}")
            detalhes_window.transient(hist_window)
            detalhes_window.grab_set()
        
        def on_tree_click(event):
            item = tree.identify_row(event.y)
            if item:
                mostrar_detalhes_final_semana(item)
        
        tree.bind("<ButtonRelease-1>", on_tree_click)
        
        def carregar():
            for item in tree.get_children():
                tree.delete(item)
            ano = int(ano_var.get()) if ano_var.get().isdigit() else None
            mes = int(mes_var.get()) if mes_var.get().isdigit() else None
            reunioes = listar_reunioes_final_semana(ano=ano, mes=mes, limite=100)
            for r in reunioes:
                data_criacao = r.get("data_criacao", "")[:19].replace("T", " ") if r.get("data_criacao") else ""
                tree.insert("", END, values=(r.get("ano", ""), r.get("mes", ""), r.get("nome_arquivo", ""), data_criacao))
        
        ttk.Button(main, text="Filtrar", command=carregar, bootstyle="primary").pack(pady=(10, 0))
        carregar()
        
        hist_window.update_idletasks()
        x = (hist_window.winfo_screenwidth() // 2) - (900 // 2)
        y = (hist_window.winfo_screenheight() // 2) - (600 // 2)
        hist_window.geometry(f"900x600+{x}+{y}")

    def historico_publicadores(self):
        # Configuração da janela
        historico_pub_window = ttk.Toplevel(self.root)
        historico_pub_window.title("Histórico de Publicadores")
        historico_pub_window.geometry("1024x768")
        
        # Container principal
        main_container = ttk.Frame(historico_pub_window, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        # Título
        title_frame = ttk.Frame(main_container)
        title_frame.pack(fill=X, pady=(0, 20))
        
        ttk.Label(
            title_frame,
            text="Histórico de Publicadores",
            font=("Helvetica", 20, "bold"),
            bootstyle="primary"
        ).pack(side=LEFT)
        
        # Frame de busca e filtros
        search_frame = ttk.Frame(main_container)
        search_frame.pack(fill=X, pady=(0, 20))
        
        # Campo de busca
        ttk.Label(
            search_frame,
            text="Buscar publicador:",
            bootstyle="secondary"
        ).pack(side=LEFT, padx=(0, 10))
        
        search_var = ttk.StringVar()
        search_entry = ttk.Entry(
            search_frame,
            textvariable=search_var,
            width=40,
            bootstyle="primary"
        )
        search_entry.pack(side=LEFT, padx=(0, 20))
        
        # Frame para a tabela
        table_frame = ttk.Frame(main_container)
        table_frame.pack(fill=BOTH, expand=YES)
        
        # Criar Treeview
        colunas = ("nome", "batizado", "ultima_parte", "detalhes")
        tree = ttk.Treeview(
            table_frame,
            columns=colunas,
            show="headings",
            bootstyle="primary",
            height=20
        )
        
        # Configurar colunas
        tree.heading("nome", text="Nome", anchor=CENTER)
        tree.heading("batizado", text="Batizado", anchor=CENTER)
        tree.heading("ultima_parte", text="Última Parte", anchor=CENTER)
        tree.heading("detalhes", text="Detalhes", anchor=CENTER)
        
        # Configurar largura das colunas
        tree.column("nome", width=300, anchor=CENTER)
        tree.column("batizado", width=100, anchor=CENTER)
        tree.column("ultima_parte", width=300, anchor=CENTER)
        tree.column("detalhes", width=100, anchor=CENTER)
        
        # Adicionar scrollbar
        scrollbar = ttk.Scrollbar(
            table_frame,
            orient="vertical",
            command=tree.yview,
            bootstyle="primary-round"
        )
        scrollbar.pack(side=RIGHT, fill=Y)
        tree.configure(yscrollcommand=scrollbar.set)
        tree.pack(fill=BOTH, expand=YES)
        
        def mostrar_detalhes_publicador(item):
            # Obter dados do publicador selecionado
            valores = tree.item(item)['values']
            nome_publicador = valores[0]
            
            # Criar janela de detalhes
            detalhes_window = ttk.Toplevel(historico_pub_window)
            detalhes_window.title(f"Detalhes do Publicador - {nome_publicador}")
            detalhes_window.geometry("800x600")
            
            # Container principal
            detalhes_container = ttk.Frame(detalhes_window, padding=20)
            detalhes_container.pack(fill=BOTH, expand=YES)
            
            # Título
            ttk.Label(
                detalhes_container,
                text=f"Histórico de Partes - {nome_publicador}",
                font=("Helvetica", 16, "bold"),
                bootstyle="primary"
            ).pack(pady=(0, 20))
            
            # Frame para a tabela de detalhes
            detalhes_table_frame = ttk.Frame(detalhes_container)
            detalhes_table_frame.pack(fill=BOTH, expand=YES)
            
            # Criar Treeview para detalhes
            detalhes_tree = ttk.Treeview(
                detalhes_table_frame,
                columns=("data", "parte"),
                show="headings",
                bootstyle="primary",
                height=15
            )
            
            # Configurar colunas
            detalhes_tree.heading("data", text="Data", anchor=CENTER)
            detalhes_tree.heading("parte", text="Parte", anchor=CENTER)
            
            # Configurar largura das colunas
            detalhes_tree.column("data", width=300, anchor=W)
            detalhes_tree.column("parte", width=400, anchor=W)
            
            # Adicionar scrollbar
            detalhes_scrollbar = ttk.Scrollbar(
                detalhes_table_frame,
                orient="vertical",
                command=detalhes_tree.yview,
                bootstyle="primary-round"
            )
            detalhes_scrollbar.pack(side=RIGHT, fill=Y)
            detalhes_tree.configure(yscrollcommand=detalhes_scrollbar.set)
            detalhes_tree.pack(fill=BOTH, expand=YES, padx=(0, 10))
            
            try:
                from database import db_ops
                # Buscar histórico do publicador
                historico = db_ops.buscar_historico_publicador(nome_publicador)
                
                if historico:
                    # Ordenar histórico por data (mais recente primeiro)
                    historico_ordenado = sorted(historico, key=lambda x: x['data'], reverse=True)
                    
                    # Inserir dados na tabela com cores alternadas
                    for i, registro in enumerate(historico_ordenado):
                        tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                        detalhes_tree.insert("", END, values=(
                            registro['data'],
                            registro['parte']
                        ), tags=(tag,))
                    
                    # Configurar cores alternadas
                    detalhes_tree.tag_configure('evenrow', background='#f0f0f0')
                    detalhes_tree.tag_configure('oddrow', background='white')
                else:
                    detalhes_tree.insert("", END, values=(
                        "Nenhum histórico encontrado",
                        ""
                    ))
            
            except Exception as e:
                logger.error(f"Erro ao carregar histórico do publicador: {str(e)}")
                detalhes_tree.insert("", END, values=(
                    f"Erro ao carregar histórico: {str(e)}",
                    ""
                ))
            
            # Botão de voltar
            ttk.Button(
                detalhes_container,
                text="Voltar",
                command=detalhes_window.destroy,
                bootstyle="secondary",
                width=15
            ).pack(side=BOTTOM, pady=(20, 0))
            
            # Centralizar a janela
            detalhes_window.update_idletasks()
            width = detalhes_window.winfo_width()
            height = detalhes_window.winfo_height()
            x = (detalhes_window.winfo_screenwidth() // 2) - (width // 2)
            y = (detalhes_window.winfo_screenheight() // 2) - (height // 2)
            detalhes_window.geometry(f"{width}x{height}+{x}+{y}")
            
            # Tornar a janela modal
            detalhes_window.transient(historico_pub_window)
            detalhes_window.grab_set()
        
        def carregar_publicadores():
            # Limpar tabela
            for item in tree.get_children():
                tree.delete(item)
            
            try:
                # Buscar publicadores
                from database import db_ops
                publicadores = db_ops.getAllPub()
                
                # Filtrar por busca
                termo_busca = search_var.get().lower()
                if termo_busca:
                    publicadores = [p for p in publicadores if termo_busca in p['nome'].lower()]
                
                if not publicadores:
                    tree.insert("", END, values=("Nenhum publicador encontrado", "", "", ""))
                    return
                
                # Ordenar por nome
                publicadores.sort(key=lambda x: x['nome'].lower())
                
                # Adicionar publicadores na tabela
                for i, pub in enumerate(publicadores):
                    tag = 'evenrow' if i % 2 == 0 else 'oddrow'
                    tree.insert("", END, values=(
                        pub['nome'],
                        "Sim" if pub.get('batizado', False) else "Não",
                        pub.get('ultima_parte', 'N/A'),
                        "↗"  # Ícone de seta diagonal para cima
                    ), tags=(tag,))
                
                # Configurar cores alternadas
                tree.tag_configure('evenrow', background='#f0f0f0')
                tree.tag_configure('oddrow', background='white')
            
            except Exception as e:
                logger.error(f"Erro ao carregar publicadores: {str(e)}")
                tree.insert("", END, values=(
                    f"Erro ao carregar publicadores: {str(e)}",
                    "",
                    "",
                    ""
                ))
        
        # Configurar evento de clique na linha
        def on_tree_click(event):
            item = tree.identify_row(event.y)
            if item:
                mostrar_detalhes_publicador(item)
        
        tree.bind("<ButtonRelease-1>", on_tree_click)
        
        # Configurar busca
        search_var.trace("w", lambda *args: carregar_publicadores())
        
        # Botão de voltar
        ttk.Button(
            main_container,
            text="Voltar",
            command=historico_pub_window.destroy,
            bootstyle="secondary",
            width=15
        ).pack(side=BOTTOM, pady=(20, 0))
        
        # Carregar dados iniciais
        carregar_publicadores()
        
        # Centralizar a janela
        historico_pub_window.update_idletasks()
        width = historico_pub_window.winfo_width()
        height = historico_pub_window.winfo_height()
        x = (historico_pub_window.winfo_screenwidth() // 2) - (width // 2)
        y = (historico_pub_window.winfo_screenheight() // 2) - (height // 2)
        historico_pub_window.geometry(f"{width}x{height}+{x}+{y}")

    def designacoes_salao(self):
        from database import (salvar_designacoes_salao, listar_designacoes_salao,
                               buscar_designacoes_salao, excluir_designacoes_salao)
        from process.designacoes_salao import DesignacoesSalao

        win = ttk.Toplevel(self.root)
        win.title("Designações Salão")
        win.geometry("1150x720")

        main = ttk.Frame(win, padding=20)
        main.pack(fill=BOTH, expand=YES)

        ttk.Label(main, text="Designações Salão",
                  font=("Helvetica", 20, "bold"), bootstyle="primary").pack(anchor=W, pady=(0, 15))

        # ── Configurações ────────────────────────────────────────────────
        cfg = ttk.LabelFrame(main, text="Configurações", padding=10)
        cfg.pack(fill=X, pady=(0, 10))

        mapa_dia = {"(nenhum)": None, "Segunda": 0, "Terça": 1, "Quarta": 2, "Quinta": 3, "Sexta": 4}
        mapa_fds = {"(nenhum)": None, "Sábado": 5, "Domingo": 6}

        ttk.Label(cfg, text="Dia Meio de Semana:").grid(row=0, column=0, padx=(0, 5), pady=5, sticky=W)
        dia_semana_var = ttk.StringVar(value="(nenhum)")
        ttk.Combobox(cfg, textvariable=dia_semana_var,
                     values=list(mapa_dia.keys()), width=12, state="readonly",
                     bootstyle="primary").grid(row=0, column=1, padx=(0, 20), pady=5, sticky=W)

        ttk.Label(cfg, text="Dia Final de Semana:").grid(row=0, column=2, padx=(0, 5), pady=5, sticky=W)
        dia_fds_var = ttk.StringVar(value="(nenhum)")
        ttk.Combobox(cfg, textvariable=dia_fds_var,
                     values=list(mapa_fds.keys()), width=12, state="readonly",
                     bootstyle="primary").grid(row=0, column=3, padx=(0, 20), pady=5, sticky=W)

        ttk.Label(cfg, text="Meses:").grid(row=0, column=4, padx=(0, 5), pady=5, sticky=W)
        meses_var = ttk.IntVar(value=3)
        ttk.Spinbox(cfg, textvariable=meses_var, from_=1, to=12, width=5,
                    bootstyle="primary").grid(row=0, column=5, padx=(0, 20), pady=5, sticky=W)

        ttk.Button(cfg, text="Gerar Automaticamente", bootstyle="success",
                   command=lambda: gerar_auto()).grid(row=0, column=6, padx=5, pady=5)
        ttk.Button(cfg, text="Carregar Dados Salvos", bootstyle="info",
                   command=lambda: carregar_salvos()).grid(row=0, column=7, padx=5, pady=5)

        # ── Tabela ────────────────────────────────────────────────────────
        tabela_frame = ttk.Frame(main)
        tabela_frame.pack(fill=BOTH, expand=YES, pady=(0, 10))

        cols = ("data", "tipo", "audio", "video", "microfone", "indicadores")
        tabela = ttk.Treeview(tabela_frame, columns=cols, show="headings",
                              bootstyle="primary", height=18,
                              displaycolumns=("data", "audio", "video", "microfone", "indicadores"))
        tabela.heading("data",        text="Data")
        tabela.heading("audio",       text="Áudio")
        tabela.heading("video",       text="Vídeo")
        tabela.heading("microfone",   text="Microfone")
        tabela.heading("indicadores", text="Indicadores")
        tabela.column("data",        width=90,  anchor=CENTER)
        tabela.column("audio",       width=160, anchor=W)
        tabela.column("video",       width=160, anchor=W)
        tabela.column("microfone",   width=220, anchor=W)
        tabela.column("indicadores", width=220, anchor=W)

        sb_v = ttk.Scrollbar(tabela_frame, orient="vertical",
                             command=tabela.yview, bootstyle="primary-round")
        sb_v.pack(side=RIGHT, fill=Y)
        tabela.configure(yscrollcommand=sb_v.set)
        sb_h = ttk.Scrollbar(tabela_frame, orient="horizontal",
                             command=tabela.xview, bootstyle="primary-round")
        sb_h.pack(side=BOTTOM, fill=X)
        tabela.configure(xscrollcommand=sb_h.set)
        tabela.pack(fill=BOTH, expand=YES)

        # ── Botões ────────────────────────────────────────────────────────
        btn_frame = ttk.Frame(main)
        btn_frame.pack(fill=X)

        ttk.Button(btn_frame, text="Salvar", bootstyle="primary",
                   command=lambda: salvar()).pack(side=LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Exportar DOCX", bootstyle="success",
                   command=lambda: exportar_docx()).pack(side=LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Excluir Mês Selecionado", bootstyle="danger",
                   command=lambda: excluir_mes()).pack(side=LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="Fechar", bootstyle="secondary",
                   command=win.destroy).pack(side=RIGHT)

        # ── Callbacks ─────────────────────────────────────────────────────
        def _popula_tabela(designacoes):
            for row in tabela.get_children():
                tabela.delete(row)
            for d in designacoes:
                tabela.insert("", "end", values=(
                    d["data"], d.get("tipo", ""),
                    d.get("audio", ""), d.get("video", ""),
                    d.get("microfone", ""), d.get("indicadores", "")
                ))

        def carregar_salvos():
            from ttkbootstrap.dialogs import Messagebox
            docs = listar_designacoes_salao()
            if not docs:
                Messagebox.show_info("Nenhuma designação salva encontrada.", "Carregar")
                return
            todas = []
            for doc in docs:
                todas.extend(doc.get("semanas", []))
            if not todas:
                Messagebox.show_info("Nenhuma designação salva encontrada.", "Carregar")
                return
            todas.sort(key=lambda x: __import__('datetime').datetime.strptime(x["data"], "%d/%m/%Y"))
            _popula_tabela(todas)

        def gerar_auto():
            idx_dia = mapa_dia.get(dia_semana_var.get())
            idx_fds = mapa_fds.get(dia_fds_var.get())
            if idx_dia is None and idx_fds is None:
                from ttkbootstrap.dialogs import Messagebox
                Messagebox.show_warning("Selecione ao menos um dia de reunião.", "Atenção")
                return
            datas = DesignacoesSalao.gerar_datas(idx_dia, idx_fds, meses_var.get())
            if not datas:
                from ttkbootstrap.dialogs import Messagebox
                Messagebox.show_warning("Nenhuma data gerada para o período.", "Atenção")
                return
            designacoes = DesignacoesSalao.gerar_designacoes(datas)
            _popula_tabela(designacoes)

        def salvar():
            import datetime as _dt
            from collections import defaultdict
            from ttkbootstrap.dialogs import Messagebox

            # Coletar dados da tabela
            por_mes = defaultdict(list)
            todas_linhas = []
            for iid in tabela.get_children():
                vals = tabela.item(iid)["values"]
                data_str, tipo, audio, video, mic, ind = (str(v) for v in vals)
                try:
                    d = _dt.datetime.strptime(data_str, "%d/%m/%Y")
                except ValueError:
                    continue
                row = {"data": data_str, "tipo": tipo, "audio": audio,
                       "video": video, "microfone": mic, "indicadores": ind}
                por_mes[(d.year, d.month)].append(row)
                todas_linhas.append(row)

            if not por_mes:
                Messagebox.show_warning("Nenhum dado para salvar.", "Atenção")
                return

            # ── Janela de pré-visualização ────────────────────────────────
            prev = ttk.Toplevel(win)
            prev.title("Validar Designações")
            prev.grab_set()
            prev_w, prev_h = 1100, 600
            px = (prev.winfo_screenwidth() // 2) - (prev_w // 2)
            py = (prev.winfo_screenheight() // 2) - (prev_h // 2)
            prev.geometry(f"{prev_w}x{prev_h}+{px}+{py}")

            ttk.Label(prev, text="Revise as designações antes de salvar:",
                      font=("Helvetica", 13, "bold"),
                      bootstyle="primary").pack(anchor=W, padx=20, pady=(15, 5))

            meses_pt = {1:"Janeiro",2:"Fevereiro",3:"Março",4:"Abril",5:"Maio",
                        6:"Junho",7:"Julho",8:"Agosto",9:"Setembro",
                        10:"Outubro",11:"Novembro",12:"Dezembro"}
            info_meses = ", ".join(
                f"{meses_pt.get(m, m)}/{a}" for a, m in sorted(por_mes.keys())
            )
            ttk.Label(prev,
                      text=f"{len(todas_linhas)} datas  |  Meses: {info_meses}",
                      bootstyle="secondary").pack(anchor=W, padx=20, pady=(0, 8))

            # Treeview de prévia
            prev_frame = ttk.Frame(prev)
            prev_frame.pack(fill=BOTH, expand=YES, padx=20)

            prev_cols = ("data", "audio", "video", "microfone", "indicadores")
            prev_tree = ttk.Treeview(prev_frame, columns=prev_cols,
                                     show="headings", bootstyle="primary", height=16)
            prev_tree.heading("data",        text="Data")
            prev_tree.heading("audio",       text="Áudio")
            prev_tree.heading("video",       text="Vídeo")
            prev_tree.heading("microfone",   text="Microfone")
            prev_tree.heading("indicadores", text="Indicadores")
            prev_tree.column("data",        width=85,  anchor=CENTER)
            prev_tree.column("audio",       width=155, anchor=W)
            prev_tree.column("video",       width=155, anchor=W)
            prev_tree.column("microfone",   width=215, anchor=W)
            prev_tree.column("indicadores", width=215, anchor=W)

            prev_sb = ttk.Scrollbar(prev_frame, orient="vertical",
                                    command=prev_tree.yview, bootstyle="primary-round")
            prev_sb.pack(side=RIGHT, fill=Y)
            prev_tree.configure(yscrollcommand=prev_sb.set)
            prev_tree.pack(fill=BOTH, expand=YES)

            for row in todas_linhas:
                prev_tree.insert("", "end", values=(
                    row["data"], row["audio"], row["video"],
                    row["microfone"], row["indicadores"]
                ))

            # Botões confirmar / cancelar
            btn_prev = ttk.Frame(prev)
            btn_prev.pack(fill=X, padx=20, pady=15)

            def confirmar_salvar():
                prev.destroy()
                for (ano, mes), semanas in por_mes.items():
                    dados = {
                        "ano": ano, "mes": mes,
                        "dia_semana": dia_semana_var.get(),
                        "dia_fds": dia_fds_var.get(),
                        "semanas": semanas
                    }
                    res = salvar_designacoes_salao(dados)
                    if not res.get("success"):
                        Messagebox.show_error(res.get("message", "Erro ao salvar."), "Erro", parent=win)
                        return
                Messagebox.show_info("Designações salvas com sucesso!", "Sucesso", parent=win)

            ttk.Button(btn_prev, text="Confirmar e Salvar", bootstyle="success",
                       command=confirmar_salvar).pack(side=LEFT, padx=(0, 10))
            ttk.Button(btn_prev, text="Cancelar", bootstyle="secondary",
                       command=prev.destroy).pack(side=LEFT)

        def excluir_mes():
            import datetime as _dt
            from ttkbootstrap.dialogs import Messagebox
            sel = tabela.selection()
            if not sel:
                Messagebox.show_warning("Selecione uma linha para identificar o mês.", "Atenção", parent=win)
                return
            vals = tabela.item(sel[0])["values"]
            data_str = str(vals[0])
            try:
                d = _dt.datetime.strptime(data_str, "%d/%m/%Y")
            except ValueError:
                Messagebox.show_warning("Data inválida na linha selecionada.", "Atenção", parent=win)
                return
            confirm = Messagebox.yesno(
                f"Excluir todas as designações de {d.month:02d}/{d.year}?\nEsta ação remove o histórico dos irmãos.",
                "Confirmar Exclusão",
                parent=win
            )
            if confirm == "Yes":
                res = excluir_designacoes_salao(d.year, d.month)
                if res.get("success"):
                    for iid in tabela.get_children():
                        v = tabela.item(iid)["values"]
                        try:
                            rd = _dt.datetime.strptime(str(v[0]), "%d/%m/%Y")
                            if rd.year == d.year and rd.month == d.month:
                                tabela.delete(iid)
                        except ValueError:
                            pass
                    Messagebox.show_info(res.get("message", "Excluído."), "Sucesso", parent=win)
                else:
                    Messagebox.show_error(res.get("message", "Erro ao excluir."), "Erro", parent=win)

        def exportar_docx():
            import datetime as _dt
            import os
            from tkinter import filedialog
            from ttkbootstrap.dialogs import Messagebox
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            rows = []
            for iid in tabela.get_children():
                vals = tabela.item(iid)["values"]
                rows.append({
                    "data":       str(vals[0]),
                    "audio":      str(vals[2]),
                    "video":      str(vals[3]),
                    "microfone":  str(vals[4]),
                    "indicadores":str(vals[5]),
                })

            if not rows:
                Messagebox.show_warning("Nenhum dado na tabela.", "Atenção")
                return

            caminho = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                initialfile="Designacoes_Salao.docx",
                title="Salvar DOCX"
            )
            if not caminho:
                return

            doc = Document()

            # ── Página paisagem, margens estreitas ────────────────────────
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = Cm(29.7), Cm(21.0)
            section.top_margin    = Cm(1.2)
            section.bottom_margin = Cm(1.2)
            section.left_margin   = Cm(1.5)
            section.right_margin  = Cm(1.5)

            # ── Título ────────────────────────────────────────────────────
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_title = p_title.add_run("Designações Salão")
            run_title.font.size = Pt(13)
            run_title.font.bold = True

            # ── Tabela ────────────────────────────────────────────────────
            # Colunas: Data | Áudio | Vídeo | Microfone | Indicadores
            col_widths = [Cm(2.2), Cm(4.8), Cm(4.8), Cm(6.5), Cm(6.5)]
            headers    = ["Data", "Áudio", "Vídeo", "Microfone", "Indicadores"]

            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            # Cabeçalho
            hdr = table.rows[0].cells
            for i, (txt, w) in enumerate(zip(headers, col_widths)):
                hdr[i].width = w
                p = hdr[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(txt)
                run.font.size = Pt(9)
                run.font.bold = True
                # Fundo azul claro
                tc = hdr[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'BDD7EE')
                tcPr.append(shd)

            meses_pt = {1:"JANEIRO",2:"FEVEREIRO",3:"MARÇO",4:"ABRIL",5:"MAIO",
                        6:"JUNHO",7:"JULHO",8:"AGOSTO",9:"SETEMBRO",
                        10:"OUTUBRO",11:"NOVEMBRO",12:"DEZEMBRO"}

            def _set_cell_fill(cell, hex_color):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), hex_color)
                tcPr.append(shd)

            # Linhas de dados com header de mês mesclado
            fill_data = ['FFFFFF', 'EBF3FB']
            prev_month = None
            month_color_idx = 0

            for row_data in rows:
                try:
                    d = _dt.datetime.strptime(row_data["data"], "%d/%m/%Y")
                    cur_month = (d.year, d.month)
                except ValueError:
                    cur_month = None

                if cur_month != prev_month:
                    month_color_idx = 1 - month_color_idx
                    prev_month = cur_month

                    # ── Linha de cabeçalho do mês (span 5 cols) ──────────
                    if cur_month:
                        mes_label = f"{meses_pt.get(cur_month[1], '')}  —  {cur_month[0]}"
                        hdr_row = table.add_row().cells
                        # Mescla todas as células
                        merged = hdr_row[0].merge(hdr_row[1]).merge(hdr_row[2]).merge(hdr_row[3]).merge(hdr_row[4])
                        p_hdr = merged.paragraphs[0]
                        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_hdr = p_hdr.add_run(mes_label)
                        run_hdr.font.size = Pt(9)
                        run_hdr.font.bold = True
                        run_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        _set_cell_fill(merged, '2F5496')

                row_cells = table.add_row().cells
                values_row = [
                    row_data["data"], row_data["audio"], row_data["video"],
                    row_data["microfone"], row_data["indicadores"]
                ]
                for i, (val, w) in enumerate(zip(values_row, col_widths)):
                    row_cells[i].width = w
                    p = row_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(val)
                    run.font.size = Pt(8)
                    fill = fill_data[month_color_idx]
                    if fill != 'FFFFFF':
                        _set_cell_fill(row_cells[i], fill)

            doc.save(caminho)
            if Messagebox.yesno(f"DOCX salvo em:\n{caminho}\n\nAbrir arquivo?", "Exportado") == "Yes":
                os.startfile(caminho)

        def on_double_click(event):
            item = tabela.identify_row(event.y)
            col  = tabela.identify_column(event.x)
            if not item or col == "#1":
                return
            try:
                disp_idx = int(col.replace("#", "")) - 1  # 0-based display index
                if disp_idx <= 0 or disp_idx > 4:
                    return
            except ValueError:
                return
            # displaycolumns=("data","audio","video","microfone","indicadores") omite "tipo" (values[1])
            # display idx → values idx: 0→0, 1→2, 2→3, 3→4, 4→5
            disp_to_val = {0: 0, 1: 2, 2: 3, 3: 4, 4: 5}
            col_labels   = {1: "Áudio", 2: "Vídeo", 3: "Microfone", 4: "Indicadores"}
            values_idx = disp_to_val[disp_idx]
            current = tabela.item(item)["values"]
            popup = ttk.Toplevel(win)
            popup.title(f"Editar — {col_labels.get(disp_idx, '')}")
            popup.resizable(False, False)
            popup.grab_set()
            entry_var = ttk.StringVar(value=str(current[values_idx]))
            ttk.Entry(popup, textvariable=entry_var, width=40,
                      bootstyle="primary").pack(padx=15, pady=(15, 5))
            if disp_idx in (3, 4):
                ttk.Label(popup, text='Formato: "Nome1 / Nome2"',
                          bootstyle="secondary", font=("Helvetica", 9)).pack()
            def confirmar(event=None):
                vals = list(tabela.item(item)["values"])
                vals[values_idx] = entry_var.get()
                tabela.item(item, values=vals)
                popup.destroy()
            ttk.Button(popup, text="OK", command=confirmar,
                       bootstyle="primary").pack(pady=(8, 15))
            popup.bind("<Return>", confirmar)

        tabela.bind("<Double-1>", on_double_click)

        # Carregar dados salvos automaticamente ao abrir
        carregar_salvos()

        # Centralizar janela
        win.update_idletasks()
        x = (win.winfo_screenwidth() // 2) - (1150 // 2)
        y = (win.winfo_screenheight() // 2) - (720 // 2)
        win.geometry(f"1150x720+{x}+{y}")

    def dashboards(self):
        # Criar janela de dashboards
        dashboards_window = ttk.Toplevel(self.root)
        dashboards_window.title("Dashboards")
        dashboards_window.geometry("1000x600")
        
        # Container principal
        main_container = ttk.Frame(dashboards_window, padding=20)
        main_container.pack(fill=BOTH, expand=YES)
        
        # Título
        title_frame = ttk.Frame(main_container)
        title_frame.pack(fill=X, pady=(0, 20))
        
        ttk.Label(
            title_frame,
            text="Dashboards",
            font=("Helvetica", 20, "bold"),
            bootstyle="primary"
        ).pack(side=LEFT)
        
        # Menu de navegação
        menu_frame = ttk.Frame(main_container)
        menu_frame.pack(fill=X, pady=(0, 20))
        
        # Variável para controlar o dashboard atual
        current_dashboard = ttk.StringVar(value="participacoes")
        
        # Botões do menu
        ttk.Button(
            menu_frame,
            text="Participações por Publicador",
            command=lambda: current_dashboard.set("participacoes"),
            bootstyle="primary-outline",
            width=25
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            menu_frame,
            text="Participações por Reunião",
            command=lambda: current_dashboard.set("reunioes"),
            bootstyle="primary-outline",
            width=25
        ).pack(side=LEFT, padx=5)

        ttk.Button(
            menu_frame,
            text="Designações Salão",
            command=lambda: current_dashboard.set("salao"),
            bootstyle="info-outline",
            width=20
        ).pack(side=LEFT, padx=5)

        # Frame para filtro de parte (apenas para dashboard de participações)
        filtro_frame = ttk.Frame(main_container)
        # Não empacotar inicialmente, será mostrado apenas quando necessário
        
        # Lista de partes disponíveis
        partes_disponiveis = [
            "Todas as Partes",
            "Todas as Partes Menos Oração",
            "Todas as Partes Menos Final de Semana",
            "Presidente",
            "Oração Inicial",
            "Tesouro",
            "Joias Espirituais",
            "Leitura da Bíblia",
            "Escola - Primeira Parte",
            "Escola - Segunda Parte",
            "Escola - Terceira Parte",
            "Escola - Quarta Parte",
            "Nossa Vida Cristã - Primeira Parte",
            "Nossa Vida Cristã - Segunda Parte",
            "Estudo de Congregação",
            "Oração Final"
        ]
        
        # Variável para controlar a parte selecionada
        parte_selecionada = ttk.StringVar(value="Todas as Partes")
        
        # Label e combobox para filtro de parte
        ttk.Label(
            filtro_frame,
            text="Filtrar por Parte:",
            bootstyle="secondary"
        ).pack(side=LEFT, padx=(0, 10))
        
        parte_combo = ttk.Combobox(
            filtro_frame,
            textvariable=parte_selecionada,
            values=partes_disponiveis,
            width=30,
            state="readonly",
            bootstyle="primary"
        )
        parte_combo.pack(side=LEFT, padx=(0, 10))
        
        # Frame para os gráficos
        graphs_frame = ttk.Frame(main_container)
        graphs_frame.pack(fill=BOTH, expand=YES)
        
        # Importar matplotlib
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        import matplotlib
        matplotlib.use('TkAgg')
        
        def atualizar_grafico(*args):
            # Mostrar/ocultar filtro de parte baseado no dashboard selecionado
            if current_dashboard.get() == "participacoes":
                filtro_frame.pack(fill=X, pady=(0, 10), before=graphs_frame)
            else:
                filtro_frame.pack_forget()
            
            # Limpar frame de gráficos
            for widget in graphs_frame.winfo_children():
                widget.destroy()

            # Dashboard de Designações Salão (usa Treeview, não matplotlib)
            if current_dashboard.get() == "salao":
                filtro_frame.pack_forget()
                from database import contar_designacoes_salao_por_publicador
                dados = contar_designacoes_salao_por_publicador()
                dados_sorted = sorted(dados.items(), key=lambda x: x[1]['total'], reverse=True)
                cols = ("nome", "audio", "video", "microfone", "indicadores", "total")
                tree_salao = ttk.Treeview(graphs_frame, columns=cols, show="headings",
                                          bootstyle="primary", height=18)
                tree_salao.heading("nome", text="Nome")
                tree_salao.heading("audio", text="Áudio")
                tree_salao.heading("video", text="Vídeo")
                tree_salao.heading("microfone", text="Microfone")
                tree_salao.heading("indicadores", text="Indicadores")
                tree_salao.heading("total", text="Total")
                tree_salao.column("nome", width=200, anchor=W)
                for c in ("audio", "video", "microfone", "indicadores", "total"):
                    tree_salao.column(c, width=90, anchor=CENTER)
                for nome, counts in dados_sorted:
                    tree_salao.insert("", "end", values=(
                        nome, counts['audio'], counts['video'],
                        counts['microfone'], counts['indicadores'], counts['total']
                    ))
                sb_salao = ttk.Scrollbar(graphs_frame, orient="vertical",
                                         command=tree_salao.yview, bootstyle="primary-round")
                sb_salao.pack(side=RIGHT, fill=Y)
                tree_salao.configure(yscrollcommand=sb_salao.set)
                tree_salao.pack(fill=BOTH, expand=YES)
                return

            # Criar nova figura
            fig, ax = plt.subplots(figsize=(10, 6))
            
            if current_dashboard.get() == "participacoes":
                # Obter filtro de parte selecionado
                parte_filtro = parte_selecionada.get()
                
                # Tratar opções especiais
                if parte_filtro == "Todas as Partes":
                    parte_para_busca = None
                elif parte_filtro == "Todas as Partes Menos Oração":
                    parte_para_busca = "__EXCLUIR_ORACOES__"
                elif parte_filtro == "Todas as Partes Menos Final de Semana":
                    parte_para_busca = "__EXCLUIR_FINAL_SEMANA__"
                else:
                    parte_para_busca = parte_filtro
                
                # Obter dados usando o novo método
                db = DatabaseOperations()
                participacoes_dict = db.contar_participacoes_por_parte(parte=parte_para_busca)
                
                if not participacoes_dict:
                    ax.text(0.5, 0.5, "Nenhuma participação encontrada", 
                            ha='center', va='center', fontsize=14, color='gray',
                            transform=ax.transAxes)
                    plt.tight_layout()
                    canvas = FigureCanvasTkAgg(fig, master=graphs_frame)
                    canvas.draw()
                    canvas.get_tk_widget().pack(fill=BOTH, expand=YES)
                    return
                
                # Limitar aos top 20 publicadores para melhor visualização
                top_publicadores = list(participacoes_dict.items())[:20]
                
                # Separar dados
                nomes = [d[0] for d in top_publicadores]
                participacoes = [d[1] for d in top_publicadores]
                
                # Criar gráfico de barras
                bars = ax.barh(nomes, participacoes)
                
                # Adicionar valores nas barras
                for i, (bar, valor) in enumerate(zip(bars, participacoes)):
                    ax.text(valor, i, f' {valor}', va='center', fontsize=9)
                
                # Configurar o gráfico
                titulo = f'Participações por Publicador'
                if parte_filtro != "Todas as Partes":
                    titulo += f' - {parte_filtro}'
                ax.set_title(titulo, fontsize=12, fontweight='bold')
                ax.set_xlabel('Número de Participações')
                ax.set_ylabel('Publicadores')
                ax.invert_yaxis()  # Inverter para mostrar o maior no topo
                plt.tight_layout()
            
            else:  # participacoes por reuniao
                # Obter dados das reuniões usando o novo método
                db = DatabaseOperations()
                total_reunioes, participacoes_unicas = db.contar_participacoes_unicas_por_reuniao()
                
                # Se não houver dados, mostrar mensagem
                if not participacoes_unicas:
                    ax.text(0.5, 0.5, "Nenhum dado de participação encontrado", 
                            ha='center', va='center', fontsize=14, color='gray',
                            transform=ax.transAxes)
                    plt.tight_layout()
                    canvas = FigureCanvasTkAgg(fig, master=graphs_frame)
                    canvas.draw()
                    canvas.get_tk_widget().pack(fill=BOTH, expand=YES)
                    return
                
                # Limitar aos top 20 publicadores para melhor visualização
                top_publicadores = list(participacoes_unicas.items())[:20]
                
                # Separar dados
                nomes = [d[0] for d in top_publicadores]
                participacoes = [d[1] for d in top_publicadores]
                
                # Criar gráfico de barras horizontal
                bars = ax.barh(nomes, participacoes)
                
                # Adicionar valores nas barras
                for i, (bar, valor) in enumerate(zip(bars, participacoes)):
                    ax.text(valor, i, f' {valor}', va='center', fontsize=9)
                
                # Adicionar texto com total de reuniões no topo do gráfico
                ax.text(0.02, 0.98, f'Total de Reuniões: {total_reunioes}', 
                        transform=ax.transAxes, fontsize=12, fontweight='bold',
                        verticalalignment='top', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
                
                # Configurar o gráfico
                ax.set_title('Participações Únicas por Reunião - Top Publicadores', 
                            fontsize=12, fontweight='bold')
                ax.set_xlabel('Número de Reuniões Participadas')
                ax.set_ylabel('Publicadores')
                ax.invert_yaxis()  # Inverter para mostrar o maior no topo
                plt.tight_layout()
            
            # Ajustar layout
            plt.tight_layout()
            
            # Criar canvas para exibir o gráfico
            canvas = FigureCanvasTkAgg(fig, master=graphs_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=BOTH, expand=YES)
        
        # Configurar trace para atualizar o gráfico quando mudar o dashboard ou a parte
        current_dashboard.trace_add("write", atualizar_grafico)
        parte_selecionada.trace_add("write", atualizar_grafico)
        
        # Carregar o gráfico inicial
        atualizar_grafico()
        
        # Botão de voltar
        ttk.Button(
            main_container,
            text="Voltar",
            command=dashboards_window.destroy,
            bootstyle="secondary",
            width=15
        ).pack(side=BOTTOM, pady=(20, 0))

    def abrir_configuracoes(self):
        import sys
        from pathlib import Path
        from util.backup import restore_database, _project_root

        dialog = ttk.Toplevel(self.root)
        dialog.title("Configurações")
        dialog.transient(self.root)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding=20)
        frame.pack(fill=BOTH, expand=YES)

        ttk.Label(frame, text="⚙ Configurações", font=("Helvetica", 16, "bold"), bootstyle="primary").pack(anchor=W, pady=(0, 15))

        # --- Restaurar Banco ---
        restore_frame = ttk.LabelFrame(frame, text="Restaurar Banco", padding=15)
        restore_frame.pack(fill=X, pady=(0, 10))

        ttk.Label(restore_frame, text="Data do backup:").pack(anchor=W)

        backups_dir = _project_root() / "backups"
        datas = sorted(
            [d.name for d in backups_dir.iterdir() if d.is_dir()],
            reverse=True
        ) if backups_dir.exists() else []

        data_var = tk.StringVar(value=datas[0] if datas else "")
        combo = ttk.Combobox(restore_frame, textvariable=data_var, values=datas, state="readonly", width=20)
        combo.pack(anchor=W, pady=(2, 10))

        colecoes_frame = ttk.Frame(restore_frame)
        colecoes_frame.pack(fill=X)

        check_vars = {}

        def atualizar_colecoes(*_):
            for w in colecoes_frame.winfo_children():
                w.destroy()
            check_vars.clear()
            data = data_var.get()
            if not data:
                return
            pasta = backups_dir / data
            arquivos = sorted(pasta.glob("*.json"))
            for arq in arquivos:
                nome = arq.stem
                var = tk.BooleanVar(value=True)
                check_vars[nome] = var
                ttk.Checkbutton(colecoes_frame, text=nome, variable=var, bootstyle="primary").pack(anchor=W)

        combo.bind("<<ComboboxSelected>>", atualizar_colecoes)
        atualizar_colecoes()

        def fazer_restore():
            selecionadas = [nome for nome, var in check_vars.items() if var.get()]
            logger.info(f"Restore solicitado: coleções={selecionadas}")
            if not selecionadas:
                Messagebox.show_warning("Selecione ao menos uma coleção.", "Atenção", parent=dialog)
                return
            data = data_var.get()
            confirmar = Messagebox.yesno(
                f"Restaurar {len(selecionadas)} coleção(ões) do backup {data}?\n\nDados atuais serão substituídos.",
                "Confirmar Restauração",
                parent=dialog
            )
            logger.info(f"Confirmação restore: {confirmar!r}")
            if confirmar not in ("Yes", "Sim"):
                return
            try:
                resultados = restore_database(data, selecionadas)
                logger.info(f"Restore concluído: {resultados}")
                linhas = [f"✓ {col}: {res['count']} registros ({res['status']})" for col, res in resultados.items()]
                Messagebox.show_info("\n".join(linhas), "Restauração Concluída", parent=dialog)
            except Exception as e:
                logger.error(f"Erro no restore: {e}")
                Messagebox.show_error(f"Erro ao restaurar:\n{e}", "Erro", parent=dialog)

        ttk.Button(
            restore_frame,
            text="Restaurar",
            bootstyle="danger",
            command=fazer_restore,
            width=15
        ).pack(anchor=W, pady=(12, 0))

        # Centralizar dialog
        dialog.update_idletasks()
        w, h = 420, 480
        x = self.root.winfo_x() + (self.root.winfo_width() - w) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - h) // 2
        dialog.geometry(f"{w}x{h}+{x}+{y}")

if __name__ == "__main__":
    try:
        # Configurar logging — arquivo gravado em %APPDATA%\JW Mural\
        import sys as _sys
        if getattr(_sys, 'frozen', False):
            _log_dir = os.path.join(os.environ.get('APPDATA', os.path.expanduser('~')), 'JW Mural')
        else:
            _log_dir = os.path.dirname(os.path.abspath(__file__))
        os.makedirs(_log_dir, exist_ok=True)
        _log_file = os.path.join(_log_dir, 'jw_mural.log')
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(_log_file, encoding='utf-8'),
                logging.StreamHandler(),
            ]
        )
        
        # Criar janela principal com tema
        root = ttk.Window(themename="litera")
        root.title("JW Mural")
        
        # Corrigir texto dos botões no Windows (evitar "só traço colorido" por fonte inválida)
        try:
            style = root.style
            font_btn = ("Segoe UI", 10)
            style.configure("TButton", font=font_btn)
            for name in ("primary", "secondary", "success", "info", "warning", "danger"):
                for suffix in ("", "-outline", "-link", "-round", "-round-toggle"):
                    try:
                        style.configure(f"{name}{suffix}.TButton", font=font_btn)
                    except tk.TclError:
                        pass
        except Exception:
            pass
        
        # Configurar geometria da janela principal
        window_width = 1024
        window_height = 768
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # Criar aplicação
        app = ModernApp(root)
        
        # Forçar atualização da janela principal
        root.update_idletasks()
        root.deiconify()
        root.update()
        
        # Função para iniciar as verificações após a janela principal estar visível
        def start_checks():
            if not initialize_application(root):
                Messagebox.show_error(
                    "Não foi possível inicializar o sistema. Verifique os logs para mais detalhes.",
                    "Erro de Inicialização"
                )
                root.quit()
                sys.exit(1)
        
        # Agendar a inicialização para acontecer após a janela principal estar pronta
        root.after(100, start_checks)
        
        # Iniciar loop principal
        root.mainloop()
        
    except Exception as e:
        logger.error(f"Erro ao iniciar a aplicação: {str(e)}")
        Messagebox.show_error(
            f"Ocorreu um erro ao iniciar a aplicação:\n{str(e)}",
            "Erro Fatal"
        )
        sys.exit(1)
