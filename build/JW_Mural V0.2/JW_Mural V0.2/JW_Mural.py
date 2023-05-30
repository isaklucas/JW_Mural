import requests
from bs4 import BeautifulSoup
import docx
import os
import subprocess






class MyApp:
        def run(urlEv, nomeEv, idiomaEv):
                
                # Variaveis digitadas pelo Usuario
                idioma = idiomaEv
                filename = nomeEv
                urlEnviada = urlEv
                
                # Pega a Semana que está na URL
                urlSplit = urlEnviada.split("/")
                semanas = urlSplit[len(urlSplit) - 1]

                # Variaveis utilizadas para os For e montar a Url novamente
                tamanhoURL = len(urlSplit)
                x = int(semanas)
                y = x+5
                base_url = ''

                for v in range(0, tamanhoURL-1) :
                        if v == 0 :
                                base_url = base_url  + urlSplit[v]
                        else : base_url = base_url + "/" + urlSplit[v]


                # Seleciona o Template de acordo com Idioma escolhido
                if idioma == "pt" :
                        template_filename = "Templates/Template_PT.docx"
                
                # Nome do arquivo para gravaçao do documento
                nomeArquivo = filename + ".docx"


                # Lendo o modelo do Word Selecionado
                doc = docx.Document(template_filename) 

                # Variavel que susbtitui no Word
                pagina = 'p'
                
                while x < y:

                        url = base_url + "/" + str(x)
                        print("Criando semana: " + str(x))
                
                        x = x + 1
                        # Fazendo a solicitação HTTP
                        response = requests.get(url)

                        # Verificando se a solicitação foi bem-sucedida
                        if response.status_code == 200:
                        
                                # Analisando o HTML com BeautifulSoup
                                soup = BeautifulSoup(response.text, "html.parser")
                                
                                # Procurando a div com a classe "todayitem"
                                pag = soup.find("div", class_="todayItem")
                                
                                # Extraindo o texto da div
                                text = pag.get_text()  
                        
                                verificaSemana = soup.find(id="p1")
                                
                                if verificaSemana != None:
                                        semana = soup.find(id="p1").get_text() 
                                        livro = soup.find(id="p2").get_text()
                                        canticoInicial = soup.find(id="p3").get_text()
                                        tesouroTemp = soup.find(id="p6").get_text().split('(')
                                        perguntaJoias = soup.find(id="p8").get_text()
                                        leituraSemana = soup.find(id="p12").get_text().split(')')
                                        canticoMeio = soup.find(id="p18").get_text()
                                        nvcP1temp = soup.find(id="p19").get_text().split('(')
                                        nvcP2temp = soup.find(id="p20").get_text().split('(')
                                        estudo = soup.find(id="p21").get_text()
                                        canticoFinal = soup.find(id="p23")

                                        nvcP1 = nvcP1temp[0]
                                        nvcP2 = nvcP2temp[0]
                                        tesouro = tesouroTemp[0]  

                                        if canticoFinal != None and nvcP2 != "\nSua resposta\n\n":
                                                canticoFinal = soup.find(id="p23").get_text()        
                                                # Localizando os campos de substituição no modelo
                                                for paragraph in doc.paragraphs:
                                                        for run in paragraph.runs:
                                                                if str(pagina) + "01" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "01", semana) 
                                                                if str(pagina) + "02" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "02", livro)
                                                                if str(pagina) + "03" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "03", canticoInicial)
                                                                if str(pagina) + "06" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "06", tesouro)
                                                                if str(pagina) + "08" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "08", perguntaJoias)
                                                                if str(pagina) + "12" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "12", leituraSemana[1] + ")")
                                                                if str(pagina) + "18" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "18", canticoMeio)
                                                                if str(pagina) + "19" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "19", nvcP1)
                                                                if str(pagina) + "20" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "20", nvcP2)
                                                                if str(pagina) + "21" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "21", estudo)
                                                                if str(pagina) + "23" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "23", canticoFinal)


                                        if canticoFinal == None :
                                                estudo = soup.find(id="p20").get_text()
                                                canticoFinal = soup.find(id="p22").get_text()
                                                # Localizando os campos de substituição no modelo
                                                for paragraph in doc.paragraphs:
                                                        for run in paragraph.runs:
                                                                if str(pagina) + "01" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "01", semana) 
                                                                if str(pagina) + "02" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "02", livro)
                                                                if str(pagina) + "03" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "03", canticoInicial)
                                                                if str(pagina) + "06" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "06", tesouro)
                                                                if str(pagina) + "08" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "08", perguntaJoias)
                                                                if str(pagina) + "12" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "12" , leituraSemana[1] + ")")
                                                                if str(pagina) + "18" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "18" , canticoMeio)
                                                                if str(pagina) + "19" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "19" , nvcP1)
                                                                if str(pagina) + "20" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "20" , "Nao possue essa parte")
                                                                if str(pagina) + "21" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "21" , estudo)
                                                                if str(pagina) + "23" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "23" , canticoFinal)
                                
                                        
                                        if nvcP2 == "\nSua resposta\n\n":

                                                nvcP2 = soup.find(id="p21").get_text()
                                                estudo = soup.find(id="p22").get_text()
                                                canticoFinal = soup.find(id="p24").get_text()
                                                # Localizando os campos de substituição no modelo
                                                for paragraph in doc.paragraphs:
                                                        for run in paragraph.runs:
                                                                if str(pagina) + "01" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "01", semana) 
                                                                if str(pagina) + "02" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "02", livro)
                                                                if str(pagina) + "03" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "03", canticoInicial)
                                                                if str(pagina) + "06" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "06", tesouro)
                                                                if str(pagina) + "08" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "08", perguntaJoias)
                                                                if str(pagina) + "12" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "12" , leituraSemana[1] + ")")
                                                                if str(pagina) + "18" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "18" , canticoMeio)
                                                                if str(pagina) + "19" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "19" , nvcP1)
                                                                if str(pagina) + "20" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "20" , "Nao possue essa parte")
                                                                if str(pagina) + "21" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "21" , estudo)
                                                                if str(pagina) + "23" in run.text:
                                                                        run.text = run.text.replace(str(pagina) + "23" , canticoFinal)          
                                else :
                                        y += 1
                                        pagina = chr(ord(pagina) -1) 
                                
                                pagina = chr(ord(pagina) + 1)

                
                if not os.path.exists("documentosCriados"):
                        os.makedirs("documentosCriados")

                doc.save("documentosCriados/" + nomeArquivo)
                # Exibindo mensagem de sucesso
                print("Programaçao Criada com Sucesso na pasta DocumentosCriados com Nome:", str(nomeArquivo))
                
                
                subprocess.run(["start", "documentosCriados/" + nomeArquivo], shell=True)
                


